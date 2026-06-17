"""
Azure Certification Notes Generator
Creates a comprehensive Word document for AZ-104 and AZ-305 exam preparation.
Author: Hasan Raza (hasansyscrowd@gmail.com)
MS ID: MS1101321878
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import datetime

doc = Document()

# ============================================================
# STYLES & FORMATTING
# ============================================================
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)

# Heading styles
for i in range(1, 5):
    heading_style = doc.styles[f'Heading {i}']
    heading_style.font.name = 'Calibri'
    heading_style.font.color.rgb = RGBColor(0x00, 0x47, 0x8A)

doc.styles['Heading 1'].font.size = Pt(24)
doc.styles['Heading 2'].font.size = Pt(18)
doc.styles['Heading 3'].font.size = Pt(14)
doc.styles['Heading 4'].font.size = Pt(12)

# Helper functions
def add_table(doc, headers, rows, col_widths=None):
    """Add a formatted table to the document."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="00478A"/>')
        cell._tc.get_or_add_tcPr().append(shading)
    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, cell_text in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(cell_text)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph()
    return table

def add_bullet_list(doc, items, bold_prefix=False):
    """Add a bullet list to the document."""
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        if bold_prefix and ':' in item:
            parts = item.split(':', 1)
            run = p.add_run(parts[0] + ':')
            run.bold = True
            run.font.size = Pt(11)
            run2 = p.add_run(parts[1])
            run2.font.size = Pt(11)
        else:
            run = p.add_run(item)
            run.font.size = Pt(11)

def add_key_point(doc, text):
    """Add a highlighted key point."""
    p = doc.add_paragraph()
    run = p.add_run('★ ')
    run.bold = True
    run.font.color.rgb = RGBColor(0xE8, 0x7C, 0x00)
    run = p.add_run(text)
    run.font.size = Pt(11)

def add_assessment_box(doc, questions_answers):
    """Add assessment Q&A in a formatted way."""
    p = doc.add_paragraph()
    run = p.add_run('Assessment Answers')
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x00, 0x78, 0xD4)
    for qa in questions_answers:
        p = doc.add_paragraph()
        run = p.add_run(f'Q: {qa[0]}')
        run.bold = True
        run.font.size = Pt(10)
        p = doc.add_paragraph()
        run = p.add_run(f'A: {qa[1]}')
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x00, 0x7A, 0x33)

def section_break(doc):
    """Add a section break."""
    doc.add_page_break()

# ============================================================
# COVER PAGE
# ============================================================
for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('MICROSOFT AZURE')
run.bold = True
run.font.size = Pt(36)
run.font.color.rgb = RGBColor(0x00, 0x47, 0x8A)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Certification Study Guide')
run.font.size = Pt(24)
run.font.color.rgb = RGBColor(0x00, 0x78, 0xD4)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('AZ-104: Microsoft Azure Administrator')
run.bold = True
run.font.size = Pt(16)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('AZ-305: Azure Solutions Architect Expert')
run.bold = True
run.font.size = Pt(16)

for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Prepared by: Hasan Raza')
run.font.size = Pt(14)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Email: hasansyscrowd@gmail.com')
run.font.size = Pt(12)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Microsoft ID: MS1101321878')
run.font.size = Pt(12)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(f'Date: {datetime.datetime.now().strftime("%B %d, %Y")}')
run.font.size = Pt(12)

section_break(doc)

# ============================================================
# TABLE OF CONTENTS (Manual)
# ============================================================
doc.add_heading('Table of Contents', level=1)
doc.add_paragraph()

toc_items = [
    'PART 1: AZ-305 - Azure Solutions Architect Expert',
    '  Section A: Prerequisites (6 Modules)',
    '    1. Core Architectural Components',
    '    2. Azure Compute Services',
    '    3. Azure Storage Services',
    '    4. Identity, Access & Security',
    '    5. Cloud Adoption Framework',
    '    6. Well-Architected Framework',
    '  Section B: Design Modules',
    '    7. Design Governance',
    '    8. Design Authentication & Authorization',
    '',
    'PART 2: AZ-104 - Microsoft Azure Administrator',
    '  Section A: Prerequisites (2 Modules)',
    '    1. Azure Cloud Shell',
    '    2. ARM Templates',
    '  Section B: Manage Identities & Governance (5 Modules)',
    '    3. Microsoft Entra ID',
    '    4. Create & Manage Identities',
    '    5. Azure Policy',
    '    6. Azure RBAC',
    '    7. Microsoft Entra SSPR',
    '  Section C: Implement & Manage Storage (4 Modules)',
    '    8. Configure Storage Accounts',
    '    9. Configure Azure Blob Storage',
    '    10. Configure Azure Storage Security',
    '    11. Configure Azure Files',
    '  Section D: Configure & Manage Virtual Networks (8 Modules)',
    '    12. Configure Virtual Networks',
    '    13. Configure NSGs',
    '    14. Host Domain on Azure DNS',
    '    15. Configure VNet Peering',
    '    16. Manage Traffic Flow with Routes',
    '    17. Azure Load Balancer',
    '    18. Azure Application Gateway',
    '    19. Azure Network Watcher',
    '  Section E: Deploy & Manage Compute Resources (5 Modules)',
    '    20. Introduction to Azure VMs',
    '    21. Configure VM Availability',
    '    22. Configure Azure App Service Plans',
    '    23. Configure Azure App Service',
    '    24. Configure Azure Container Instances',
    '  Section F: Monitor & Back up Azure Resources (3 Modules)',
    '    25. Introduction to Azure Backup',
    '    26. Protect VMs with Azure Backup',
    '    27. Monitor Azure VMs with Azure Monitor',
]

for item in toc_items:
    p = doc.add_paragraph()
    if item.startswith('PART'):
        run = p.add_run(item)
        run.bold = True
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(0x00, 0x47, 0x8A)
    elif item.startswith('  Section'):
        run = p.add_run(item)
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x00, 0x78, 0xD4)
    elif item.strip():
        run = p.add_run(item)
        run.font.size = Pt(11)

section_break(doc)

# ============================================================
# PART 1: AZ-305 PREREQUISITES
# ============================================================
doc.add_heading('PART 1: AZ-305 - Azure Solutions Architect Expert', level=1)
doc.add_paragraph()
doc.add_heading('Section A: Prerequisites', level=2)

# --- Module 1: Core Architectural Components ---
doc.add_heading('Module 1: Core Architectural Components', level=2)

doc.add_heading('Azure Regions', level=3)
p = doc.add_paragraph()
p.add_run('Azure has more global regions than any other cloud provider. Regions are geographical areas containing one or more datacenters. A region is a set of datacenters deployed within a latency-defined perimeter and connected through a dedicated regional low-latency network.').font.size = Pt(11)

add_bullet_list(doc, [
    'Azure has 60+ regions worldwide',
    'Some regions have special capabilities (e.g., US Gov Virginia for government)',
    'Not all Azure services are available in every region',
    'Some services are global and not tied to a specific region (e.g., Azure AD, Traffic Manager, Azure Front Door)',
])

doc.add_heading('Availability Zones', level=3)
p = doc.add_paragraph()
p.add_run('Availability Zones are physically separate locations within an Azure region. Each zone has independent power, cooling, and networking. There are a minimum of three zones in all enabled regions.').font.size = Pt(11)

add_bullet_list(doc, [
    'Zonal services: Pin resource to a specific zone (e.g., VMs, managed disks, IP addresses)',
    'Zone-redundant services: Platform replicates automatically across zones (e.g., zone-redundant storage, SQL Database)',
    'Non-regional services: Always available from Azure global infrastructure, resilient to zone/region outages',
])

doc.add_heading('Region Pairs', level=3)
add_bullet_list(doc, [
    'Most Azure regions are paired with another region within the same geography at least 300 miles away',
    'Allows replication of resources across geography to reduce interruptions',
    'If one region has an outage, services automatically failover to the paired region',
    'Planned Azure updates are rolled out to paired regions one region at a time',
    'Data continues to reside within the same geography (except Brazil South)',
])

doc.add_heading('Azure Management Hierarchy', level=3)
p = doc.add_paragraph()
p.add_run('Azure provides four levels of management scope:').font.size = Pt(11)

add_table(doc, ['Level', 'Description'], [
    ['Management Groups', 'Containers for managing access, policies, and compliance across multiple subscriptions. All subscriptions in a management group inherit conditions applied to it.'],
    ['Subscriptions', 'Logically associate user accounts with resources. Each subscription has limits/quotas on resources. Organizations can use subscriptions to manage costs and resources.'],
    ['Resource Groups', 'Logical containers for deploying and managing Azure resources. Every resource must be in one and only one resource group.'],
    ['Resources', 'Instances of services you create (VMs, storage, databases, etc.)'],
])

add_key_point(doc, 'Management Group hierarchy: Root > Management Groups > Subscriptions > Resource Groups > Resources')
add_key_point(doc, 'Maximum 6 levels of depth for Management Groups (not including root or subscription level)')
add_key_point(doc, 'RBAC is NOT enabled by default on management groups')

# --- Module 2: Azure Compute Services ---
section_break(doc)
doc.add_heading('Module 2: Azure Compute Services', level=2)

doc.add_heading('Azure Virtual Machines', level=3)
p = doc.add_paragraph()
p.add_run('Azure VMs provide IaaS in the form of a virtualized server. You have total control over the OS and software.').font.size = Pt(11)

doc.add_heading('VM Size Families', level=4)
add_table(doc, ['Series', 'Use Case', 'Description'], [
    ['B-series', 'Dev/Test', 'Burstable, economical. Good for workloads that don\'t need full CPU continuously'],
    ['D-series', 'General Purpose', 'Enterprise-grade for most production workloads. Balanced CPU-to-memory ratio'],
    ['E-series', 'Memory Optimized', 'High memory-to-core ratio. Databases, in-memory analytics, SAP HANA'],
    ['F-series', 'Compute Optimized', 'High CPU-to-memory ratio. Gaming, batch processing, analytics'],
    ['M-series', 'Memory Optimized', 'Largest memory VMs in Azure. SAP HANA, SQL Server big data'],
    ['L-series', 'Storage Optimized', 'High disk throughput and I/O. Big data, NoSQL, data warehousing'],
    ['N-series', 'GPU Enabled', 'GPU-enabled. Deep learning, AI, graphics-intensive, video editing'],
])

doc.add_heading('VM Scale Sets', level=3)
add_bullet_list(doc, [
    'Deploy and manage a set of identical, auto-scaling VMs',
    'Number of instances can automatically increase or decrease based on demand or schedule',
    'Supports up to 1,000 VM instances (600 with custom images)',
    'Provides high availability and application resiliency',
])

doc.add_heading('Azure Virtual Desktop', level=3)
add_bullet_list(doc, [
    'Desktop and application virtualization service in the cloud',
    'Full Windows desktop experience - supports multi-session Windows 10/11',
    'Accessible from any device with a Remote Desktop client',
    'Centralized management through Azure portal',
])

doc.add_heading('Azure Containers', level=3)
add_table(doc, ['Service', 'Description'], [
    ['Azure Container Instances (ACI)', 'Fastest and simplest way to run a container in Azure. PaaS - no VM management. Suitable for isolated containers, simple apps, task automation.'],
    ['Azure Container Apps', 'Serverless container service. Scale based on HTTP traffic, events, or custom metrics. Supports microservices with Dapr integration.'],
    ['Azure Kubernetes Service (AKS)', 'Fully managed Kubernetes orchestration service. For complex, multi-container workloads that need full Kubernetes capabilities.'],
])

doc.add_heading('Azure Functions', level=3)
add_bullet_list(doc, [
    'Serverless compute service - event-driven, runs code on demand',
    'Stateless (default): Behave as if restarted every time they respond to an event',
    'Stateful (Durable Functions): Context is passed through the function to track prior activity',
    'Scale automatically based on demand',
    'You are charged only for the CPU time used while your function runs',
])

doc.add_heading('Azure App Service', level=3)
add_bullet_list(doc, [
    'PaaS for hosting web apps, REST APIs, and mobile backends',
    'Supports .NET, .NET Core, Java, Ruby, Node.js, PHP, Python',
    'Built-in auto-scale and load balancing',
    'CI/CD integration with Azure DevOps, GitHub, Bitbucket, Docker Hub',
    'Supports Windows and Linux environments',
])

# --- Module 3: Azure Storage Services ---
section_break(doc)
doc.add_heading('Module 3: Azure Storage Services', level=2)

doc.add_heading('Storage Account Types', level=3)
add_table(doc, ['Account Type', 'Supported Services', 'Performance Tiers'], [
    ['Standard general-purpose v2', 'Blob, Queue, Table, Azure Files', 'Standard (HDD)'],
    ['Premium block blobs', 'Blob (incl. Data Lake)', 'Premium (SSD)'],
    ['Premium file shares', 'Azure Files', 'Premium (SSD)'],
    ['Premium page blobs', 'Page blobs only', 'Premium (SSD)'],
])

doc.add_heading('Storage Redundancy Options', level=3)
add_table(doc, ['Redundancy', 'Copies', 'Description'], [
    ['LRS (Locally Redundant)', '3 copies', 'Replicates data 3 times within a single datacenter. Lowest cost. 11 nines durability.'],
    ['ZRS (Zone Redundant)', '3 copies', 'Replicates across 3 availability zones in the primary region. 12 nines durability.'],
    ['GRS (Geo Redundant)', '6 copies', 'LRS in primary + LRS in secondary region (300+ miles away). 16 nines durability.'],
    ['GZRS (Geo-Zone Redundant)', '6 copies', 'ZRS in primary + LRS in secondary region. 16 nines durability.'],
    ['RA-GRS (Read-Access Geo)', '6 copies', 'GRS + read access to secondary region.'],
    ['RA-GZRS (Read-Access Geo-Zone)', '6 copies', 'GZRS + read access to secondary region. Highest availability and durability.'],
])

doc.add_heading('Azure Blob Storage Access Tiers', level=3)
add_table(doc, ['Tier', 'Min Retention', 'Use Case', 'Storage Cost', 'Access Cost'], [
    ['Hot', 'None', 'Frequently accessed data', 'Highest', 'Lowest'],
    ['Cool', '30 days', 'Infrequently accessed, short-term backup', 'Lower', 'Higher'],
    ['Cold', '90 days', 'Rarely accessed, stored for 90+ days', 'Even lower', 'Even higher'],
    ['Archive', '180 days', 'Rarely accessed, flexible latency (hours)', 'Lowest', 'Highest'],
])

add_key_point(doc, 'Archive tier: Offline - must rehydrate (change tier) before accessing data')
add_key_point(doc, 'Only Hot and Cool can be set at account level. Cold and Archive are blob-level only.')

doc.add_heading('Azure Files', level=3)
add_bullet_list(doc, [
    'Fully managed file shares in the cloud accessible via SMB, NFS, or HTTP (REST)',
    'Can replace or supplement on-premises file servers',
    'Azure File Sync: Cache Azure file shares on on-premises Windows servers',
])

doc.add_heading('Azure Queues', level=3)
add_bullet_list(doc, [
    'Service for storing large numbers of messages',
    'Messages can be up to 64 KB each',
    'Queue can contain millions of messages up to total capacity of storage account',
])

doc.add_heading('Azure Disks (Managed)', level=3)
add_bullet_list(doc, [
    'Block-level storage volumes managed by Azure, used with VMs',
    'Types: Ultra disks, Premium SSD v2, Premium SSD, Standard SSD, Standard HDD',
])

doc.add_heading('Azure Migrate', level=3)
add_bullet_list(doc, [
    'Central hub for migrating to Azure',
    'Tools: Azure Migrate: Discovery and Assessment, Server Migration',
    'Supports servers, databases, web apps, virtual desktops, large datasets',
])

doc.add_heading('Azure Data Box', level=3)
add_bullet_list(doc, [
    'Physical migration service for large amounts of data (>40 TB)',
    'Microsoft ships proprietary storage device (max 80 TB capacity)',
    'Data Box Disk: Up to 35 TB via SSD disks',
    'Data Box Heavy: Up to 1 PB capacity',
])

# --- Module 4: Identity, Access & Security ---
section_break(doc)
doc.add_heading('Module 4: Identity, Access & Security', level=2)

doc.add_heading('Microsoft Entra ID (formerly Azure AD)', level=3)
add_bullet_list(doc, [
    'Cloud-based identity and access management service',
    'Single Sign-On (SSO) across thousands of cloud applications',
    'Multi-Factor Authentication (MFA): Something you know + something you have + something you are',
    'Passwordless authentication: Windows Hello for Business, Microsoft Authenticator, FIDO2 security keys',
])

doc.add_heading('External Identities', level=3)
add_table(doc, ['Type', 'Description'], [
    ['B2B Collaboration', 'Guest users access your apps with their own identity (work/school/social). Invitation and redemption flow.'],
    ['B2B Direct Connect', 'Mutual trust with another Entra organization. Supports Teams shared channels.'],
    ['Azure AD B2C / External ID', 'Publish consumer-facing apps. Custom-branded identity for customers. Separate tenant.'],
])

doc.add_heading('Conditional Access', level=3)
add_bullet_list(doc, [
    'IF-THEN policies: If a user wants to access a resource, THEN they must complete an action',
    'Requires Microsoft Entra ID P1 or P2 license',
    'Signals: User/group, IP location, device, application, risk detection',
    'Decisions: Block access, Grant access (may require MFA, compliant device, etc.)',
    'Report-only mode: Test policies before enforcing',
    'What If tool: Evaluate which policies apply to specific scenarios',
])

doc.add_heading('Azure RBAC (Role-Based Access Control)', level=3)
add_bullet_list(doc, [
    'Authorization system built on Azure Resource Manager for fine-grained access management',
    'Uses an allow model: When assigned a role, RBAC allows you to perform actions',
    'Applied at a scope: Management group > Subscription > Resource group > Resource',
    'Effective permissions = sum of all role assignments',
])

add_table(doc, ['Built-in Role', 'Description'], [
    ['Owner', 'Full access to all resources + can delegate access to others'],
    ['Contributor', 'Can create and manage all resources but cannot grant access'],
    ['Reader', 'Can view existing resources only'],
    ['User Access Administrator', 'Can manage user access to Azure resources'],
])

doc.add_heading('Zero Trust Model', level=3)
add_bullet_list(doc, [
    'Verify explicitly: Always authenticate and authorize based on all available data points',
    'Use least privilege access: Limit user access with Just-In-Time and Just-Enough-Access (JIT/JEA)',
    'Assume breach: Minimize blast radius, segment access, verify end-to-end encryption',
])

doc.add_heading('Defense-in-Depth', level=3)
p = doc.add_paragraph()
p.add_run('7 layers (from outer to inner):').font.size = Pt(11)
add_bullet_list(doc, [
    'Physical security: Protect building access to hardware in datacenter',
    'Identity & access: Control access to infrastructure, use SSO and MFA, audit events',
    'Perimeter: DDoS protection, perimeter firewalls',
    'Network: Limit communication between resources, deny by default, restrict inbound internet, secure connectivity',
    'Compute: Secure access to VMs, endpoint protection, keep systems patched',
    'Application: Ensure apps are secure and free of vulnerabilities, store secrets securely',
    'Data: Controls to manage access to data (business/customer)',
])

doc.add_heading('Microsoft Defender for Cloud', level=3)
add_bullet_list(doc, [
    'Cloud Security Posture Management (CSPM) and Cloud Workload Protection Platform (CWPP)',
    'Assesses environment continuously and provides security score',
    'Works with Azure, on-premises, and multi-cloud (AWS, GCP)',
    'Provides security alerts and advanced threat protection',
])

doc.add_heading('Azure Key Vault', level=3)
add_bullet_list(doc, [
    'Centralized cloud service for storing secrets, keys, and certificates',
    'Secrets management: Securely store and control access to tokens, passwords, certificates, API keys, etc.',
    'Key management: Encryption keys used for data encryption',
    'Certificate management: Provision, manage, and deploy public and private TLS/SSL certificates',
])

# --- Module 5: Cloud Adoption Framework ---
section_break(doc)
doc.add_heading('Module 5: Cloud Adoption Framework (CAF)', level=2)

p = doc.add_paragraph()
p.add_run('The Microsoft Cloud Adoption Framework for Azure is guidance designed to help organizations create and implement strategies for cloud adoption. It consists of 8 methodologies:').font.size = Pt(11)

add_table(doc, ['Methodology', 'Description'], [
    ['Strategy', 'Define business justification and expected outcomes of adoption'],
    ['Plan', 'Align actionable adoption plans to business outcomes'],
    ['Ready', 'Prepare the cloud environment for planned changes (landing zones)'],
    ['Migrate', 'Migrate existing workloads to cloud'],
    ['Modernize', 'Improve and modernize deployed workloads'],
    ['Cloud-native', 'Build new cloud-native or hybrid solutions'],
    ['Govern', 'Govern the environment and workloads'],
    ['Manage', 'Operations management for cloud and hybrid solutions'],
])

add_key_point(doc, 'Strategy methodology comes first - define WHY you are moving to the cloud')
add_key_point(doc, 'Security is integrated across all methodologies')

# --- Module 6: Well-Architected Framework ---
section_break(doc)
doc.add_heading('Module 6: Azure Well-Architected Framework', level=2)

p = doc.add_paragraph()
p.add_run('The Azure Well-Architected Framework consists of five pillars of architectural excellence:').font.size = Pt(11)

add_table(doc, ['Pillar', 'Description', 'Key Principles'], [
    ['Reliability', 'Ability of a system to recover from failures and continue to function', 'High availability, disaster recovery, backup and restore, testing resilience'],
    ['Security', 'Protecting applications and data from threats', 'Defense-in-depth, identity management, encryption, network security'],
    ['Cost Optimization', 'Managing costs to maximize value delivered', 'Plan and estimate, provision with optimization, use monitoring and analytics'],
    ['Operational Excellence', 'Operations processes that keep a system running in production', 'DevOps practices, automation, monitoring, design for operations'],
    ['Performance Efficiency', 'Ability of a system to adapt to changes in load', 'Autoscaling, optimize network performance, optimize storage, identify bottlenecks'],
])

add_key_point(doc, 'Use the Azure Advisor tool for personalized recommendations aligned with these pillars')
add_key_point(doc, 'Trade-offs between pillars are common - e.g., higher security may increase cost')

# ============================================================
# PART 1 SECTION B: AZ-305 DESIGN MODULES
# ============================================================
section_break(doc)
doc.add_heading('Section B: Design Modules', level=2)

# --- Module 7: Design Governance ---
doc.add_heading('Module 7: Design Governance', level=2)

doc.add_heading('Design for Management Groups', level=3)
add_bullet_list(doc, [
    'Containers above subscriptions for organizing resources',
    'Maximum 6 levels of depth (not including root or subscription level)',
    'Flat hierarchy of 3-4 levels recommended',
    'RBAC is NOT enabled by default on management groups',
    'All subscriptions within a management group trust the same Entra ID tenant',
    'Policies and RBAC applied at management group level are inherited by all child resources',
    'A management group tree can support up to six levels of depth',
    'Each directory has a single top-level root management group',
    '10,000 management groups can be supported in a single directory',
    'New subscriptions are placed under the root management group by default',
], bold_prefix=True)

doc.add_heading('Tailwind Traders - Management Group Example', level=4)
p = doc.add_paragraph()
p.add_run('The Tailwind Traders company has the following departmental structure mapped to management groups:').font.size = Pt(11)
add_table(doc, ['Management Group', 'Sub-Groups', 'Subscriptions'], [
    ['Root Management Group', 'Sales, Corporate, IT', 'N/A'],
    ['Sales', 'West, East', 'Sales subscriptions for each region'],
    ['Corporate', 'HR, Legal', 'Corporate subscriptions per department'],
    ['IT', 'Research, Dev, Production', 'Separate subscriptions per environment'],
])

doc.add_heading('Management Group Design Considerations', level=4)
add_bullet_list(doc, [
    'Governance alignment: Align management groups with your organization\'s governance requirements',
    'Flat hierarchy: Keep hierarchy to 3-4 levels for manageability',
    'Top-level management group: Apply broad policies at root, use lower levels for exceptions',
    'Organizational/Departmental: Group by department or business unit (Sales, IT, Finance)',
    'Geographical: Group by region for region-specific compliance or data residency',
    'Production management group: Apply strict policies for production workloads',
    'Sandbox management group: Relaxed policies for development and experimentation',
    'Isolated sensitive info: Separate management group for highly regulated workloads (PCI, HIPAA)',
])

doc.add_heading('Design for Subscriptions', level=3)
add_bullet_list(doc, [
    'Subscriptions are units of management, billing, and scale',
    'Logical container that associates user accounts and their created resources',
    'Virtual networks cannot be shared across subscriptions',
    'Each subscription has limits and quotas for resources',
], bold_prefix=True)

doc.add_heading('Subscription Types', level=4)
add_table(doc, ['Type', 'Description'], [
    ['Enterprise Agreement (EA)', 'Negotiated pricing for organizations. 3-year commitment.'],
    ['Pay-As-You-Go', 'Monthly billing based on actual usage. No commitment.'],
    ['MSDN/Visual Studio', 'Monthly Azure credits for development/testing.'],
    ['Free Trial', '12 months of popular free services + $200 credit for 30 days.'],
])

doc.add_heading('Subscription Design Considerations (Tailwind Traders)', level=4)
add_bullet_list(doc, [
    'Treat subscriptions as a democratized unit of management: Each team/project can have its own subscription',
    'Group subscriptions together under management groups for unified policy and RBAC',
    'Consider a shared services subscription: Common resources (DNS, networking, monitoring) in a central subscription',
    'Be aware of subscription scale limits: E.g., max 250 storage accounts per region per subscription',
    'Simplify admin management: Minimize number of subscription owners',
    'Apply Azure Policies at subscription boundary: Enforce standards per environment (dev/prod)',
    'Consider network topologies: VNet peering or VPN gateways needed for cross-subscription communication',
    'Conduct owner access reviews: Use Privileged Identity Management (PIM) for quarterly or biannual reviews',
])

doc.add_heading('Design for Resource Groups', level=3)
add_bullet_list(doc, [
    'Cannot be nested inside other resource groups',
    'Cannot be renamed',
    'Each resource can only be in one resource group',
    'Resources can be moved between resource groups',
    'Resource group metadata is stored in the region you specify for the group',
    'Deleting a resource group deletes ALL resources inside it',
    'Resource groups can contain resources from different regions',
    'Apply resource locks (Delete or ReadOnly) to prevent accidental changes',
])

doc.add_heading('Resource Group Grouping Strategies (Tailwind Traders)', level=4)
add_table(doc, ['Strategy', 'Example', 'Best For'], [
    ['By Resource Type', 'SQL-RG, WEB-RG, STORAGE-RG', 'Simple environments, easy to locate resources by type'],
    ['By Application', 'App1-RG, App2-RG', 'Group all resources for one application together'],
    ['By Department', 'Finance-RG, Marketing-RG, IT-RG', 'Billing and access control per department'],
    ['By Location', 'EastUS-RG, WestEurope-RG', 'Region-specific compliance requirements'],
    ['By Billing', 'Project1-RG, Project2-RG', 'Track costs per project or cost center'],
    ['Combination', 'App1-Prod-EastUS-RG', 'Comprehensive naming combining app, env, and region'],
])

doc.add_heading('Resource Group Design Considerations', level=4)
add_bullet_list(doc, [
    'Life cycle: Group resources with same lifecycle together (created and deleted together)',
    'Administrative overhead: Balance between too many and too few resource groups',
    'Access control: Apply RBAC at resource group level to manage who can access what',
    'Compliance: Group resources by compliance requirements (PCI, HIPAA, etc.)',
    'Tailwind Traders example: App1 has SQL DB + App Service + Storage in App1-RG, App2 has separate App2-RG',
])

doc.add_heading('Design for Resource Tags', level=3)
add_bullet_list(doc, [
    'Name-value pairs applied to Azure resources for logical organization',
    'Tags are NOT inherited from resource group to resources (use Azure Policy to enforce)',
    'Maximum 50 tags per resource/resource group',
    'Tag name: 512 characters max, Tag value: 256 characters max',
])

add_table(doc, ['Category', 'Example Tags', 'Purpose'], [
    ['Functional', 'app, tier, env', 'Categorize resources by function or purpose'],
    ['Classification', 'confidentiality, SLA', 'Classify by policies applied'],
    ['Accounting', 'department, costCenter, billing', 'Associate resources with billing groups'],
    ['Partnership', 'owner, stakeholder, contact', 'Identify people related to resources'],
    ['Purpose', 'businessProcess, compliance, DR', 'Align resources to business functions'],
])

doc.add_heading('IT-Aligned vs Business-Aligned Tagging', level=4)
add_table(doc, ['Approach', 'Tag Examples', 'Tailwind Traders Example'], [
    ['IT-Aligned', 'Workload, Environment (dev/prod), Datacenter, Application, Approver', 'Tag printers in IT department by workload=printing, env=production'],
    ['Business-Aligned', 'CostCenter, BusinessUnit, BusinessCriticality, Revenue impact, SLA', 'Tag marketing resources by costcenter=marketing, criticality=high'],
])
add_key_point(doc, 'IT-aligned tagging: Best for tracking workloads, environments, and infrastructure decisions')
add_key_point(doc, 'Business-aligned tagging: Best for accounting, business ownership, and cost tracking')
add_key_point(doc, 'Tags are NOT inherited from parent (resource group/subscription). Use Azure Policy with "Inherit a tag" effect to auto-apply tags.')

doc.add_heading('Design for Azure Policy', level=3)
add_bullet_list(doc, [
    'Policies are inherited down the hierarchy (management group > subscription > resource group)',
    'Azure Policy evaluates resources every 24 hours (can also be triggered on demand)',
    'Policy vs RBAC: Policy enforces properties on resources; RBAC controls user actions',
    'Azure Arc enables extending Azure Policy to on-premises and multi-cloud resources',
    'Azure Policy can audit, deny, modify, or deploy resources',
    'Use initiatives to group multiple policy definitions',
])

doc.add_heading('Design for RBAC', level=3)
add_bullet_list(doc, [
    'Allow model: Effective permissions = sum of all role assignments',
    'Best practice: Assign roles to groups, not individual users',
    'Built-in roles: Owner, Contributor, Reader, User Access Administrator',
    'Custom roles can be created for specific needs',
    'Deny assignments can be created via Azure Blueprints',
    'Use PIM for just-in-time privileged access',
])

doc.add_heading('Design for Azure Landing Zones', level=3)
add_bullet_list(doc, [
    'An architecture and reference implementation for Azure environments',
    'Two types: Platform landing zones and Application landing zones',
    'Platform landing zones: Shared services (identity, connectivity, management)',
    'Application landing zones: Individual workloads and applications',
    'IaC accelerator available with Bicep and Terraform',
    'Design areas: Billing, Identity, Network, Resource Organization, Governance, Security, Management, Platform Automation',
])

# --- Module 8: Design Authentication & Authorization ---
section_break(doc)
doc.add_heading('Module 8: Design Authentication & Authorization', level=2)

doc.add_heading('Design for Identity & Access Management (IAM)', level=3)
p = doc.add_paragraph()
p.add_run('Four pillars of IAM:').font.size = Pt(11)
add_bullet_list(doc, [
    'Unified identity management: Single identity system across cloud and on-premises',
    'Seamless user experience: SSO, reduce password fatigue, consistent experience',
    'Secure adaptive access: Strong authentication, risk-based conditional access',
    'Simplified identity governance: Control access for all users across all environments',
])

doc.add_heading('Design for Microsoft Entra ID', level=3)
add_bullet_list(doc, [
    'Cloud-only vs Hybrid deployment',
    'Entra Connect and Cloud Sync for hybrid scenarios',
    'Phishing-resistant authentication methods (FIDO2, Windows Hello)',
    'SSO for seamless access to cloud and on-premises applications',
    'Custom domains for branding (replace onmicrosoft.com)',
])

doc.add_heading('Design for Microsoft Entra B2B', level=3)
add_bullet_list(doc, [
    'Guest users access your applications with their own identity',
    'Guest users manage their own identity (password resets, MFA)',
    'Invitation and redemption process for guest access',
    'Guest users represented in your directory as UserType = Guest',
    'Control what guest users can see and do with guest access restrictions',
])

doc.add_heading('Design for Azure AD B2C / External ID', level=3)
add_bullet_list(doc, [
    'Customer identity and access management (CIAM)',
    'Separate tenant specifically for consumer/customer accounts',
    'Custom-branded identity experience for external users',
    'Azure AD B2C no longer available for new customers (since May 2025)',
    'Existing B2C supported until May 2030',
    'New customers should use Microsoft Entra External ID',
    'Supports social identity providers (Google, Facebook, etc.)',
])

doc.add_heading('Design for Conditional Access', level=3)
add_bullet_list(doc, [
    'Requires Microsoft Entra ID P1 or P2 license',
    'IF-THEN policies based on signals (user, location, device, app, risk)',
    'Named locations: Define trusted IP ranges or countries',
    'Report-only mode: Evaluate policy impact without enforcing',
    'What If tool: Simulate policies against specific scenarios',
    'Optimization Agent: Suggests policy improvements',
    'Common policies: Require MFA for admins, block legacy authentication, require compliant devices',
])

doc.add_heading('Design for Identity Protection', level=3)
add_table(doc, ['Risk Type', 'Examples', 'Detection'], [
    ['User Risk', 'Leaked credentials, Threat intelligence', 'Detected offline, indicates user account may be compromised'],
    ['Sign-in Risk', 'Anonymous IP address, Atypical travel, Password spray, Malicious IP, Anomalous token, Verified threat actor IP', 'Detected in real-time or offline, indicates suspicious sign-in attempt'],
])

doc.add_heading('Design for Access Reviews', level=3)
add_bullet_list(doc, [
    'Periodic review of user access to ensure only authorized users have access',
    '3 reviewer types: Resource owners, Delegates (managers), End users (self-review)',
    'Can be scheduled (weekly, monthly, quarterly, annually)',
    'Auto-apply results option to automatically remove access',
    'Useful for: Guest access cleanup, role assignment review, group membership review',
])

doc.add_heading('Design for Service Principals', level=3)
add_bullet_list(doc, [
    'Identity for applications in Entra ID',
    '3 types: Application (most common), Managed Identity, Legacy',
    'Application object is 1:many relationship with service principals',
    'App object exists in home tenant, service principal exists in each tenant where app is used',
    'Used for: Non-interactive authentication, automation, CI/CD pipelines',
])

doc.add_heading('Design for Managed Identities', level=3)
add_table(doc, ['Type', 'Lifecycle', 'Sharing', 'Use Case'], [
    ['System-assigned', 'Tied to resource (created/deleted with resource)', 'Cannot be shared', 'Single resource workloads, need independent identities'],
    ['User-assigned', 'Standalone (independent lifecycle)', 'Can be shared across resources', 'Multiple resources needing same identity, pre-authorization scenarios'],
])

add_key_point(doc, 'Managed identities eliminate the need to manage credentials in code')
add_key_point(doc, 'No additional cost for managed identities')

doc.add_heading('Design for Azure Key Vault', level=3)
add_bullet_list(doc, [
    'Stores 3 types: Secrets, Keys, Certificates',
    'Two tiers: Standard (software encryption) and Premium (HSM-backed keys)',
    'Soft delete: Enabled by default, retains deleted vaults/objects for recovery',
    'Purge protection: Prevents permanent deletion during retention period',
    'Access control via RBAC or Key Vault access policies',
    'Network restrictions: VNet service endpoints, private endpoints, firewall rules',
    'Logging and monitoring with Azure Monitor',
])

add_assessment_box(doc, [
    ('Design Governance: Which Azure AD feature supports requesting access to groups or application access packages?', 'Entitlement management'),
    ('Which Azure RBAC role allows managing user access but not resources?', 'User Access Administrator'),
    ('Design Auth: What provides single set of credentials across cloud and on-premises?', 'Hybrid identity with Entra Connect'),
])

# ============================================================
# PART 2: AZ-104 PREREQUISITES
# ============================================================
section_break(doc)
doc.add_heading('PART 2: AZ-104 - Microsoft Azure Administrator', level=1)
doc.add_paragraph()
doc.add_heading('Section A: Prerequisites', level=2)

# --- Module 1: Azure Cloud Shell ---
doc.add_heading('Module 1: Azure Cloud Shell', level=2)

add_bullet_list(doc, [
    'Browser-based shell experience for managing Azure resources',
    'Choose between Bash or PowerShell',
    'Accessible via: portal.azure.com, shell.azure.com, Azure mobile app, VS Code terminal',
    'CloudDrive: Persistent storage backed by Azure Files share',
    'Automatically authenticated with your Azure credentials',
    '20-minute idle timeout (session disconnects after 20 min of inactivity)',
    'Requires a resource group, storage account, and Azure Files share',
    'CloudDrive directory persists across sessions',
    'Azure CLI uses az commands (e.g., az vm create)',
    'Azure PowerShell uses cmdlets (e.g., New-AzVM)',
])

doc.add_heading('Azure CLI vs Azure PowerShell', level=3)
add_table(doc, ['Feature', 'Azure CLI', 'Azure PowerShell'], [
    ['Syntax', 'az <command>', 'Verb-Noun cmdlets (e.g., Get-AzVM)'],
    ['Platform', 'Cross-platform (Windows, macOS, Linux)', 'Cross-platform (PowerShell Core)'],
    ['Output', 'JSON by default', 'Objects by default'],
    ['Learning curve', 'Easier for Linux/Bash users', 'Easier for Windows admins'],
])

# --- Module 2: ARM Templates ---
doc.add_heading('Module 2: Azure Resource Manager (ARM) Templates', level=2)

add_bullet_list(doc, [
    'Infrastructure as Code (IaC): JSON files defining resources to deploy',
    'Declarative syntax: Define WHAT to deploy, not HOW',
    'Idempotent: Same template produces same result every time',
    'Submitted to ARM, which orchestrates deployment of resources in parallel',
])

doc.add_heading('ARM Template Structure', level=3)
add_table(doc, ['Element', 'Required', 'Description'], [
    ['$schema', 'Yes', 'Location of JSON schema file that describes template language version'],
    ['contentVersion', 'Yes', 'Version of the template (e.g., 1.0.0.0)'],
    ['parameters', 'No', 'Values provided during deployment to customize resources. Max 256 parameters.'],
    ['variables', 'No', 'Values used as JSON fragments to simplify template language expressions'],
    ['functions', 'No', 'User-defined functions available within the template'],
    ['resources', 'Yes', 'Resource types deployed or updated in the resource group'],
    ['outputs', 'No', 'Values returned after deployment'],
])

add_key_point(doc, 'ARM Templates support: Orchestration, what-if preview, validation, and modular deployment (linked/nested templates)')
add_key_point(doc, 'Bicep is the recommended alternative to JSON ARM templates - cleaner syntax, same capabilities')

# ============================================================
# PART 2 SECTION B: MANAGE IDENTITIES & GOVERNANCE
# ============================================================
section_break(doc)
doc.add_heading('Section B: Manage Identities & Governance', level=2)

# --- Module 3: Microsoft Entra ID ---
doc.add_heading('Module 3: Microsoft Entra ID', level=2)

doc.add_heading('What is Microsoft Entra ID?', level=3)
add_bullet_list(doc, [
    'PaaS identity service (Identity as a Service)',
    'Every Azure subscription has a trusted relationship with an Entra ID tenant',
    'Default domain: <yourorg>.onmicrosoft.com (can add custom domains)',
    'Flat structure (no OUs like on-premises AD)',
    'Uses REST API (not LDAP like traditional AD)',
    'Uses SAML, OpenID Connect, OAuth for authentication (not Kerberos/NTLM)',
])

doc.add_heading('Entra ID vs Active Directory Domain Services (AD DS)', level=3)
add_table(doc, ['Feature', 'Entra ID', 'AD DS'], [
    ['Structure', 'Flat', 'Hierarchical (OUs, domains, forests)'],
    ['Query protocol', 'REST API (HTTP/HTTPS)', 'LDAP'],
    ['Authentication', 'SAML, OpenID Connect, OAuth', 'Kerberos, NTLM'],
    ['Federation', 'Built-in', 'Requires ADFS'],
    ['Management', 'Azure portal, Microsoft 365', 'Group Policy, SCCM'],
])

doc.add_heading('Entra ID Editions', level=3)
add_table(doc, ['Edition', 'Key Features'], [
    ['Free', 'Basic identity management, SSO (up to 10 apps), MFA'],
    ['P1', 'Dynamic groups, self-service group management, Conditional Access, on-premises writeback, Microsoft Identity Manager'],
    ['P2', 'Identity Protection (risk-based policies), Privileged Identity Management (PIM), Access Reviews, Entitlement Management'],
])

doc.add_heading('Microsoft Entra Domain Services', level=3)
add_bullet_list(doc, [
    'Managed domain services: Domain join, Group Policy, LDAP, Kerberos/NTLM',
    'No need to deploy, manage, or patch domain controllers',
    'Integrates with existing Entra ID tenant',
    'Supports legacy apps that need traditional AD protocols in the cloud',
])

# --- Module 4: Create & Manage Identities ---
doc.add_heading('Module 4: Create, Configure & Manage Identities', level=2)

doc.add_heading('User Account Types', level=3)
add_table(doc, ['Type', 'Description'], [
    ['Cloud Identity', 'Created and managed entirely in Entra ID'],
    ['Directory-synced Identity', 'Synced from on-premises AD via Entra Connect'],
    ['Guest User', 'External user invited via B2B collaboration'],
])

doc.add_heading('Group Types', level=3)
add_table(doc, ['Group Type', 'Description'], [
    ['Security Group', 'Used to manage member and computer access to shared resources. Can contain users, devices, service principals, and other groups.'],
    ['Microsoft 365 Group', 'Collaboration group with shared mailbox, calendar, files, SharePoint site. Members can only be users.'],
])

doc.add_heading('Group Membership Types', level=3)
add_table(doc, ['Assignment', 'Description'], [
    ['Assigned', 'Manually add specific users as members'],
    ['Dynamic User', 'Automatically add/remove users based on attribute rules (requires P1)'],
    ['Dynamic Device', 'Automatically add/remove devices based on attribute rules (Security groups only, requires P1)'],
])

doc.add_heading('Device Identity Options', level=3)
add_table(doc, ['Option', 'Description'], [
    ['Entra Registered', 'Personal devices (BYOD). Signed in with local account, access to org resources. Supports Windows 10+, iOS, Android, macOS.'],
    ['Entra Joined', 'Organization-owned devices. Signed in with work account only. Cloud-only or hybrid. Windows 10+ and Server 2019+.'],
    ['Hybrid Entra Joined', 'Devices joined to both on-premises AD and Entra ID. Signed in with work account. Requires Entra Connect.'],
])

doc.add_heading('Other Identity Features', level=3)
add_bullet_list(doc, [
    'Group-based licensing: Assign licenses to groups instead of individual users',
    'Custom security attributes: Business-specific attributes (e.g., project name, cost center)',
    'SCIM 2.0: Protocol for automated user provisioning/deprovisioning',
    'Administrative Units: Restrict administrative scope to specific departments/regions',
])

# --- Module 5: Azure Policy ---
section_break(doc)
doc.add_heading('Module 5: Azure Policy', level=2)

doc.add_heading('Policy Resources', level=3)
add_table(doc, ['Resource', 'Description'], [
    ['Policy Definition', 'A rule that defines what to evaluate and what action to take (JSON format)'],
    ['Policy Initiative', 'A collection of policy definitions grouped for a common goal'],
    ['Policy Assignment', 'A policy/initiative applied to a specific scope'],
    ['Policy Exemption', 'Excludes a resource from policy evaluation'],
    ['Policy Attestation', 'Manual compliance evidence for manual effect policies'],
    ['Policy Remediation', 'Brings non-compliant resources into compliance'],
])

doc.add_heading('Policy Effects', level=3)
add_table(doc, ['Effect', 'Description', 'When Evaluated'], [
    ['Disabled', 'Policy is turned off', 'N/A'],
    ['Append', 'Adds fields to a resource during creation/update', 'Before deployment'],
    ['Modify', 'Adds, updates, or removes properties/tags on resource', 'Before deployment'],
    ['Deny', 'Prevents resource creation/update that violates policy', 'Before deployment'],
    ['DenyAction', 'Prevents specific actions (e.g., delete)', 'During action'],
    ['Audit', 'Creates warning event in Activity Log', 'After deployment'],
    ['AuditIfNotExists', 'Audits if a related resource does not exist', 'After deployment'],
    ['DeployIfNotExists', 'Deploys a related resource if it does not exist', 'After deployment'],
    ['Manual', 'Self-attestation of compliance', 'Manual'],
])

doc.add_heading('Policy Evaluation', level=3)
add_bullet_list(doc, [
    'Control plane enforcement (ARM operations) - not data plane by default',
    'Policies evaluate every 24 hours automatically',
    'Also triggered by: Resource creation/update, policy/initiative assignment updates, standard compliance scan',
    'Greenfield: Policies enforced on new resources from the start',
    'Brownfield: Existing resources audited and may need remediation',
    'Compliance states: Compliant, Non-compliant, Exempt, Not started, Not registered, Conflicting',
    'enforcementMode: Default (enforce) or DoNotEnforce (audit only)',
    'Safe deployment practice: Use rings to gradually roll out policies',
])

doc.add_heading('Policy Definition Structure (JSON)', level=3)
add_bullet_list(doc, [
    'displayName: Name shown in portal',
    'description: What the policy does',
    'mode: All (all resource types) or Indexed (only types that support tags and location)',
    'policyRule: IF condition THEN effect',
    'Logical operators: not, allOf (AND), anyOf (OR)',
    'Parameters: Make policies reusable with different values',
])

# --- Module 6: Azure RBAC ---
doc.add_heading('Module 6: Azure RBAC', level=2)

doc.add_heading('RBAC Fundamentals', level=3)
add_bullet_list(doc, [
    'Authorization system for managing access to Azure resources',
    'Role assignment = Security Principal + Role Definition + Scope',
    'Security principals: User, Group, Service principal, Managed identity',
    'Scope hierarchy: Management group > Subscription > Resource group > Resource',
    'Allow model: Permissions are additive (sum of all assignments)',
    'Access control (IAM) pane in Azure portal for managing role assignments',
    'Activity log shows all RBAC changes for auditing',
])

doc.add_heading('Key Built-in Roles', level=3)
add_table(doc, ['Role', 'Access Level'], [
    ['Owner', 'Full access + manage access for others'],
    ['Contributor', 'Full access but cannot manage access'],
    ['Reader', 'View all resources, no changes'],
    ['User Access Administrator', 'Manage user access only'],
])

add_key_point(doc, 'Custom roles can be created using Azure portal, PowerShell, CLI, or REST API')
add_key_point(doc, 'Always assign roles to groups instead of individual users for easier management')

# --- Module 7: Microsoft Entra SSPR ---
doc.add_heading('Module 7: Microsoft Entra SSPR (Self-Service Password Reset)', level=2)

add_bullet_list(doc, [
    'Allows users to reset their own passwords without admin intervention',
    'Requires Microsoft Entra ID P1 or P2 license (paid subscription)',
    'SSPR Scope: None (disabled) | Selected (specific group) | All (all users)',
    'Admin accounts always require 2 authentication methods (cannot be changed)',
    'Reduces IT helpdesk calls for password resets',
])

doc.add_heading('SSPR Flow (How It Works)', level=3)
p = doc.add_paragraph()
p.add_run('When a user clicks "Can\'t access your account?" the following steps occur:').font.size = Pt(11)
add_bullet_list(doc, [
    '1. Localization: Portal checks browser locale and renders SSPR page in the appropriate language',
    '2. Verification (CAPTCHA): User verifies they are not a robot',
    '3. Authentication: User enters their user ID and passes the CAPTCHA',
    '4. Password Reset: User provides a new password that meets the organization\'s password policy',
    '5. Notification: User is notified that their password has been reset',
])

doc.add_heading('Authentication Methods for SSPR', level=3)
add_table(doc, ['Method', 'Available For', 'Notes'], [
    ['Mobile app notification', 'Users and Admins', 'Push notification to Microsoft Authenticator'],
    ['Mobile app code', 'Users and Admins', 'Time-based code from Authenticator app'],
    ['Email', 'Users and Admins', 'Verification code sent to alternate email'],
    ['Mobile phone (SMS)', 'Users and Admins', 'SMS code to registered phone'],
    ['Office phone', 'Users and Admins', 'Voice call to office phone'],
    ['Security questions', 'Users ONLY', 'NOT available for admins, NOT for MFA, registration only'],
])

add_key_point(doc, 'Users must register required number of methods (1 or 2, admin configurable). Admins always need 2 methods.')

doc.add_heading('SSPR Implementation Steps', level=3)
p = doc.add_paragraph()
p.add_run('Prerequisites: Entra ID P1+ license, Authentication Policy Administrator role, non-admin test user, security group for SSPR testing').font.size = Pt(11)
add_bullet_list(doc, [
    '1. Properties: Enable SSPR for "None", "Selected" (specific group), or "All"',
    '2. Authentication methods: Choose number of methods required (1 or 2) and which methods are available',
    '3. Registration: Require users to register when signing in (Yes/No), re-confirm days (0-730, 0=never)',
    '4. Notifications: Notify users on password reset (Yes/No), notify admins when other admins reset password (Yes/No)',
    '5. Customization: Customize helpdesk link with custom email or URL',
])

doc.add_heading('SSPR Exercise Steps', level=3)
add_bullet_list(doc, [
    'Create security group "SSPRTesters" and add test user',
    'Enable SSPR for "Selected" group (SSPRTesters)',
    'Register authentication methods at: https://aka.ms/ssprsetup',
    'Test password reset at: https://aka.ms/sspr',
    'Combined registration (SSPR + MFA) at: https://mysignins.microsoft.com',
    'Verification: Email code is sent, enter code, set new password',
])

doc.add_heading('Custom Directory Branding', level=3)
add_table(doc, ['Element', 'Specification', 'Details'], [
    ['Background Image', 'PNG or JPG, 1920x1080 px', 'Max file size: 300 KB'],
    ['Company Logo', 'PNG or JPG, 32x32 px', 'Max file size: 5 KB'],
    ['Favicon', 'ICO format', 'Browser tab icon'],
    ['Page Background Color', 'Hex color code', 'Applied if background image cannot load'],
])
add_bullet_list(doc, [
    'Branding applies to: Sign-in page, SSPR page, Access Panel',
    'Test branding at: https://login.microsoft.com',
    'Supports multiple locale-specific brandings',
])

doc.add_heading('Password Writeback', level=3)
add_bullet_list(doc, [
    'Syncs cloud password reset back to on-premises Active Directory',
    'Two deployment options:',
    '  - Microsoft Entra Connect: Traditional sync agent on Windows Server',
    '  - Cloud Sync: Lightweight cloud provisioning agent, side-by-side with Entra Connect',
    'Supports: Password reset, Password change, Account unlock',
    'Requires P1 license minimum',
    'Can deploy both Entra Connect and Cloud Sync side-by-side',
])

# ============================================================
# PART 2 SECTION C: IMPLEMENT & MANAGE STORAGE
# ============================================================
section_break(doc)
doc.add_heading('Section C: Implement & Manage Storage', level=2)

# --- Module 8: Configure Storage Accounts ---
doc.add_heading('Module 8: Configure Storage Accounts', level=2)

doc.add_heading('Data Categories', level=3)
add_table(doc, ['Category', 'Examples'], [
    ['Virtual Machine Data', 'Disks (managed/unmanaged) and page blobs'],
    ['Unstructured Data', 'Blob storage, Data Lake Storage (binary and text)'],
    ['Structured Data', 'Tables, Cosmos DB, Azure SQL Database'],
])

doc.add_heading('Storage Account Types', level=3)
add_table(doc, ['Type', 'Supported Services', 'Redundancy Options', 'Usage'], [
    ['Standard GPv2', 'Blob, Queue, Table, Files', 'LRS, ZRS, GRS, RA-GRS, GZRS, RA-GZRS', 'Most storage scenarios. Default choice.'],
    ['Premium Block Blobs', 'Blob (including Data Lake)', 'LRS, ZRS', 'High transaction rates, small objects, low latency'],
    ['Premium File Shares', 'Files', 'LRS, ZRS', 'Enterprise file shares, high performance'],
    ['Premium Page Blobs', 'Page blobs only', 'LRS', 'VM disks, databases'],
])

doc.add_heading('Storage Redundancy', level=3)
add_table(doc, ['Option', 'Copies', 'Region(s)', 'Durability (nines)', 'Key Feature'], [
    ['LRS', '3', 'Single datacenter', '11 nines', 'Lowest cost, protects against server/drive failure'],
    ['ZRS', '3', '3 availability zones', '12 nines', 'Protects against datacenter failure'],
    ['GRS', '6', 'Primary + Secondary region', '16 nines', 'Protects against regional outage'],
    ['RA-GRS', '6', 'Primary + Secondary region', '16 nines', 'Read access to secondary region'],
    ['GZRS', '6', '3 AZs + Secondary region', '16 nines', 'Highest protection'],
    ['RA-GZRS', '6', '3 AZs + Secondary region', '16 nines', 'Highest protection + read access to secondary'],
])

doc.add_heading('Storage Endpoints', level=3)
add_table(doc, ['Service', 'Default Endpoint Format'], [
    ['Blob', 'https://<storage-account>.blob.core.windows.net'],
    ['Table', 'https://<storage-account>.table.core.windows.net'],
    ['Queue', 'https://<storage-account>.queue.core.windows.net'],
    ['File', 'https://<storage-account>.file.core.windows.net'],
])

doc.add_heading('Secure Storage Endpoints', level=3)
add_table(doc, ['Feature', 'Description'], [
    ['Service Endpoints', 'Extends VNet identity to storage. Traffic stays on Azure backbone. Free.'],
    ['Private Endpoints', 'Private IP address from your VNet. Traffic over private link. Most secure.'],
    ['Firewall rules', 'Restrict access to specific VNets, IP ranges, or Azure services'],
])

# --- Module 9: Configure Azure Blob Storage ---
section_break(doc)
doc.add_heading('Module 9: Configure Azure Blob Storage', level=2)

doc.add_heading('Blob Container Access Levels', level=3)
add_table(doc, ['Level', 'Description'], [
    ['Private (default)', 'No anonymous read access. Only authorized requests.'],
    ['Blob', 'Anonymous read access to blobs only (not container listing)'],
    ['Container', 'Anonymous read access to blobs + container listing'],
])

doc.add_heading('Blob Access Tiers', level=3)
add_table(doc, ['Tier', 'Min Days', 'Storage Cost', 'Access Cost', 'Use Case'], [
    ['Hot', 'None', 'Highest', 'Lowest', 'Frequently accessed data'],
    ['Cool', '30', 'Lower', 'Higher', 'Short-term backup, infrequent access'],
    ['Cold', '90', 'Even lower', 'Even higher', 'Stored 90+ days, rarely accessed'],
    ['Archive', '180', 'Lowest', 'Highest', 'Long-term backup, flexible latency (hours)'],
])

add_key_point(doc, 'Pricing rule: As tier gets cooler, storage cost DECREASES but access cost INCREASES')
add_key_point(doc, 'Archive tier is offline - must rehydrate before accessing')
add_key_point(doc, 'Early deletion fees apply if data is moved/deleted before minimum retention period')

doc.add_heading('Blob Types', level=3)
add_table(doc, ['Type', 'Max Size', 'Use Case'], [
    ['Block Blob', '190.7 TB', 'Text and binary data. Files, images, videos.'],
    ['Append Blob', '195 GB', 'Optimized for append operations. Logging.'],
    ['Page Blob', '8 TB', 'Random read/write operations. VM disks (VHDs).'],
])

doc.add_heading('Lifecycle Management', level=3)
add_bullet_list(doc, [
    'If-Then rules to automate tier transitions and data deletion',
    'Rules evaluate daily and apply to blob level',
    'Actions: tierToCool, tierToCold, tierToArchive, delete',
    'Can filter by blob name prefixes or blob index tags',
    'Available for GPv2 and Premium Block Blob accounts',
])

doc.add_heading('Object Replication', level=3)
add_bullet_list(doc, [
    'Asynchronous copy of blobs between storage accounts',
    'Requires blob versioning enabled on both source and destination',
    'Does NOT support blob snapshots',
    'Supported for Hot, Cool, and Cold tiers',
    'Source and destination can be in different regions',
])

doc.add_heading('Blob Storage Tools', level=3)
add_table(doc, ['Tool', 'Description'], [
    ['Azure Storage Explorer', 'GUI application for managing storage on desktop (Windows, macOS, Linux)'],
    ['AzCopy', 'Command-line utility for copying data to/from storage. Supports SAS tokens.'],
    ['Data Box Disk', 'Physical SSD disks for transferring up to 35 TB to Azure'],
])

# --- Module 10: Configure Azure Storage Security ---
section_break(doc)
doc.add_heading('Module 10: Configure Azure Storage Security', level=2)

doc.add_heading('Encryption', level=3)
add_bullet_list(doc, [
    'Storage Service Encryption (SSE): 256-bit AES encryption for all data at rest',
    'Encryption at rest: Automatic, enabled by default, cannot be disabled',
    'Encryption in transit: HTTPS enforced by default',
    'Customer-managed keys: Use Azure Key Vault or Managed HSM',
    'Infrastructure encryption (double encryption): Additional layer at storage service level',
])

doc.add_heading('Authorization Strategies', level=3)
add_table(doc, ['Strategy', 'Description'], [
    ['Microsoft Entra ID', 'OAuth 2.0 integration. Recommended approach. RBAC-based.'],
    ['Shared Key', 'Storage account access keys. Two keys provided for rotation.'],
    ['Shared Access Signature (SAS)', 'Delegated access with granular permissions and time limits.'],
    ['Anonymous', 'Public read access to blobs/containers (must be explicitly enabled).'],
])

doc.add_heading('Shared Access Signatures (SAS)', level=3)
add_bullet_list(doc, [
    'URI that grants restricted access to storage resources',
    'Granular control: Services, resource types, permissions, IP range, protocol, time window',
])

add_table(doc, ['SAS Type', 'Description'], [
    ['User Delegation SAS', 'Secured with Entra ID credentials + SAS token. Most secure. Blob storage only.'],
    ['Service SAS', 'Secured with storage account key. Delegates access to single service (Blob/Queue/Table/Files).'],
    ['Account SAS', 'Secured with storage account key. Delegates access across multiple services.'],
])

doc.add_heading('SAS URI Parameters', level=3)
add_table(doc, ['Parameter', 'Description'], [
    ['sv', 'Storage service version'],
    ['ss', 'Storage service (b=blob, q=queue, t=table, f=file)'],
    ['st', 'Start time'],
    ['se', 'Expiry time'],
    ['sr', 'Storage resource type'],
    ['sp', 'Permissions (r=read, w=write, d=delete, l=list, a=add, c=create, u=update, p=process)'],
    ['sip', 'IP range restriction'],
    ['spr', 'Protocol (https, https/http)'],
    ['sig', 'Signature (HMAC-SHA256)'],
])

doc.add_heading('Storage Security Best Practices', level=3)
add_bullet_list(doc, [
    'Use Entra ID authorization instead of Shared Key when possible',
    'Rotate storage account keys regularly',
    'Use User Delegation SAS (most secure type)',
    'Set SAS expiration times to minimum needed',
    'Use stored access policies for service SAS (allows revocation)',
    'Enable Defender for Storage: Malware scanning, sensitive data detection, activity-based threat detection',
    'Storage Insights: Monitor storage performance, capacity, and availability',
])

# --- Module 11: Configure Azure Files ---
section_break(doc)
doc.add_heading('Module 11: Configure Azure Files', level=2)

doc.add_heading('Azure Files Overview', level=3)
add_bullet_list(doc, [
    'Fully managed file shares in the cloud',
    'Protocols: SMB (port 445), NFS, HTTP (REST API)',
    'Can replace or supplement on-premises file servers',
    'Mount on Windows, Linux, and macOS simultaneously',
])

doc.add_heading('Storage Tiers for File Shares', level=3)
add_table(doc, ['Tier', 'Storage Type', 'Account Type', 'Redundancy'], [
    ['Premium', 'SSD', 'FileStorage', 'LRS, ZRS'],
    ['Transaction Optimized', 'HDD', 'GPv2', 'LRS, ZRS, GRS, GZRS'],
    ['Hot', 'HDD', 'GPv2', 'LRS, ZRS, GRS, GZRS'],
    ['Cool', 'HDD', 'GPv2', 'LRS, ZRS, GRS, GZRS'],
])

doc.add_heading('Authentication Methods', level=3)
add_bullet_list(doc, [
    'Identity-based: On-premises AD DS, Entra Domain Services, Entra Kerberos (hybrid identities)',
    'Storage account access key: Provides super-user access',
    'Shared Access Signature (SAS): Granular, time-limited access',
])

doc.add_heading('File Share Snapshots', level=3)
add_bullet_list(doc, [
    'Incremental: Only data changed since last snapshot is saved',
    'Read-only: Cannot be modified after creation',
    'Up to 200 snapshots per share',
    'Retained until explicitly deleted',
    'Captured at file share level, retrieval at individual file level',
])

doc.add_heading('Soft Delete', level=3)
add_bullet_list(doc, [
    'Recover accidentally deleted file shares',
    'Retention period: 1 to 365 days',
    'Enabled at storage account level (applies to all file shares in account)',
])

doc.add_heading('Azure File Sync', level=3)
add_bullet_list(doc, [
    'Centralizes file shares in Azure while keeping flexibility of on-premises file server',
    '5 components: Storage Sync Service, Sync group, Cloud endpoint, Server endpoint, Agent',
    'Cloud tiering: Frequently accessed files cached locally, infrequent files tiered to cloud',
    'Multi-site access: Azure File Sync across multiple server endpoints',
    'Up to 100 sync groups per Storage Sync Service',
    'Up to 100 server endpoints per sync group',
])

# ============================================================
# PART 2 SECTION D: CONFIGURE & MANAGE VIRTUAL NETWORKS
# ============================================================
section_break(doc)
doc.add_heading('Section D: Configure & Manage Virtual Networks', level=2)

# --- Module 12: Configure Virtual Networks ---
doc.add_heading('Module 12: Configure Virtual Networks', level=2)

doc.add_heading('Virtual Network Fundamentals', level=3)
add_bullet_list(doc, [
    'VNet provides isolation and segmentation using CIDR address blocks',
    'Address spaces must not overlap when peering VNets',
    'Subnets divide VNets into smaller segments',
    'Resources in different subnets within same VNet can communicate by default',
])

doc.add_heading('Reserved IP Addresses (5 per subnet)', level=3)
add_table(doc, ['Address', 'Purpose'], [
    ['x.x.x.0', 'Network address (identifies the subnet)'],
    ['x.x.x.1', 'Default gateway'],
    ['x.x.x.2, x.x.x.3', 'Azure DNS (mapped to VNet)'],
    ['x.x.x.255', 'Network broadcast'],
])

add_key_point(doc, 'A /29 subnet gives you only 3 usable IP addresses (8 total - 5 reserved = 3 usable)')

doc.add_heading('IP Addressing', level=3)
add_table(doc, ['Type', 'Allocation', 'Description'], [
    ['Private IP', 'Dynamic (DHCP) or Static', 'Internal communication within VNet. Not routable on internet.'],
    ['Public IP', 'Dynamic or Static', 'Internet-facing communication. Assigned to NICs, load balancers, VPN gateways, etc.'],
])

add_bullet_list(doc, [
    'Standard SKU public IP: Always static allocation, zone-redundant by default, more secure (closed to inbound by default)',
    'Basic SKU public IP: Dynamic or static, open to inbound by default (being retired)',
    'Private IP ranges: 10.0.0.0/8, 172.16.0.0/12, 192.168.0.0/16',
])

# --- Module 13: Configure NSGs ---
doc.add_heading('Module 13: Configure Network Security Groups (NSGs)', level=2)

doc.add_heading('NSG Fundamentals', level=3)
add_bullet_list(doc, [
    'Contains security rules that allow or deny network traffic',
    'Can be associated to: Subnet (max 1 NSG) or NIC (max 1 NSG)',
    'Security rules: Priority 100-4096 (lower number = higher priority)',
    'Default rules (cannot be deleted but can be overridden): DenyAllInbound, AllowVNetInBound, AllowAzureLoadBalancerInBound, AllowInternetOutBound',
])

doc.add_heading('NSG Rule Processing', level=3)
add_table(doc, ['Direction', 'Processing Order'], [
    ['Inbound', 'Subnet NSG evaluated FIRST, then NIC NSG'],
    ['Outbound', 'NIC NSG evaluated FIRST, then Subnet NSG'],
])

add_key_point(doc, 'Traffic must be allowed by BOTH NSGs (subnet and NIC) for it to flow')

doc.add_heading('Application Security Groups (ASGs)', level=3)
add_bullet_list(doc, [
    'Group VMs by workload (e.g., WebServers, DbServers) instead of explicit IPs',
    'Use ASGs as source/destination in NSG rules',
    'Simplifies rule management for workload-based security',
    'NSG augmented rules: Combine multiple IPs, ports, and service tags in a single rule',
    'Service tags: Predefined labels for Azure services (e.g., Internet, VirtualNetwork, AzureLoadBalancer, Storage, Sql)',
])

# --- Module 14: Azure DNS ---
doc.add_heading('Module 14: Host Domain on Azure DNS', level=2)

doc.add_heading('DNS Fundamentals', level=3)
add_bullet_list(doc, [
    'Azure DNS hosts DNS zones for domain name resolution',
    'Domain delegation: Update NS records at your registrar to point to Azure DNS name servers',
    'Azure DNS does NOT support purchasing domain names (use App Service domains or third-party)',
])

doc.add_heading('DNS Record Types', level=3)
add_table(doc, ['Record', 'Description'], [
    ['A', 'Maps domain to IPv4 address'],
    ['AAAA', 'Maps domain to IPv6 address'],
    ['CNAME', 'Maps domain to another domain name (alias). Cannot be used at zone apex.'],
    ['MX', 'Mail exchange - routes email'],
    ['TXT', 'Text record for verification, SPF, etc.'],
    ['NS', 'Name server for the zone'],
    ['SOA', 'Start of Authority - administrative info for the zone'],
])

doc.add_heading('Private DNS Zones', level=3)
add_bullet_list(doc, [
    'Name resolution within virtual networks without custom DNS solution',
    'Virtual network links connect private DNS zones to VNets',
    'Auto-registration: Automatically create DNS records for VMs in linked VNet',
    'Supports forward and reverse lookup',
])

doc.add_heading('Alias Records', level=3)
add_bullet_list(doc, [
    'Point directly to Azure resources (Traffic Manager, CDN, Public IP, Front Door)',
    'Automatically update when resource IP changes',
    'Prevent dangling DNS records (stale records pointing to deleted resources)',
    'Can be used at zone apex (unlike CNAME)',
    'Supported for A, AAAA, and CNAME record types',
])

# --- Module 15: VNet Peering ---
doc.add_heading('Module 15: Configure VNet Peering', level=2)

add_bullet_list(doc, [
    'Regional peering: VNets in same region',
    'Global peering: VNets in different regions',
    'Non-overlapping address spaces required',
    'Peering is nontransitive: VNet A <> VNet B and VNet B <> VNet C does NOT mean VNet A <> VNet C',
    'Gateway transit: Share VPN/ExpressRoute gateway through peered VNet (hub-spoke model)',
    'Hub-and-spoke topology: Central hub VNet with shared services, spoke VNets for workloads',
    'Peering status: Initiated > Connected (both sides must complete for traffic to flow)',
])

doc.add_heading('Extending Peering Connectivity', level=3)
add_bullet_list(doc, [
    'User-defined routes (UDR): Custom routing through Network Virtual Appliances',
    'Service chaining: Route traffic through intermediate VNets',
    'Azure Virtual Network Manager: Centrally manage VNet connectivity and security at scale',
])

# --- Module 16: Manage Traffic Flow with Routes ---
doc.add_heading('Module 16: Manage Traffic Flow with Routes', level=2)

doc.add_heading('Route Types', level=3)
add_table(doc, ['Type', 'Description'], [
    ['System Routes', 'Automatically created by Azure: VNet, Internet, None. Cannot be deleted.'],
    ['Custom Routes (UDR)', 'User-defined routes to override system routes. Created in route tables.'],
    ['BGP Routes', 'From on-premises network via VPN Gateway or ExpressRoute.'],
])

doc.add_heading('Next Hop Types', level=3)
add_table(doc, ['Next Hop', 'Description'], [
    ['Virtual Appliance', 'Route through NVA (firewall, WAN optimizer, router, IDS/IPS)'],
    ['Virtual Network Gateway', 'Route through VPN or ExpressRoute gateway'],
    ['Virtual Network', 'Route within the VNet (override default system route)'],
    ['Internet', 'Route to internet'],
    ['None', 'Drop traffic (blackhole route)'],
])

doc.add_heading('Route Priority', level=3)
add_bullet_list(doc, [
    '1st: Longest prefix match (most specific route wins)',
    '2nd: If same prefix: User-defined routes (UDR) > BGP routes > System routes',
    'Service tags can be used in UDR for simplified routing to Azure services',
])

doc.add_heading('Network Virtual Appliance (NVA)', level=3)
add_bullet_list(doc, [
    'VM that performs network functions: Firewall, WAN optimizer, Router, Load balancer, IDS/IPS',
    'Available from Azure Marketplace',
    'IP forwarding must be enabled on NVA NIC',
    'Microsegmentation: Layer 4 and Layer 7 security (OSI model)',
])

# --- Module 17: Azure Load Balancer ---
doc.add_heading('Module 17: Azure Load Balancer', level=2)

add_bullet_list(doc, [
    'Operates at OSI Layer 4 (Transport layer - TCP/UDP)',
    'Two types: Public Load Balancer (internet-facing) and Internal Load Balancer (private)',
])

doc.add_heading('Components', level=3)
add_table(doc, ['Component', 'Description'], [
    ['Front-end IP', 'Public or private IP that clients connect to'],
    ['Load balancing rules', 'Map front-end IP:port to back-end pool:port'],
    ['Back-end pool', 'VMs or VMSS instances receiving traffic'],
    ['Health probes', 'Check if back-end instances are healthy (TCP, HTTP, HTTPS)'],
    ['Inbound NAT rules', 'Forward traffic from specific front-end port to specific back-end VM'],
    ['HA ports', 'Protocol=All, Port=0. Load balance all TCP/UDP flows on all ports.'],
    ['Outbound rules', 'Configure outbound SNAT for back-end pool VMs'],
])

doc.add_heading('Session Persistence', level=3)
add_table(doc, ['Mode', 'Hash', 'Description'], [
    ['None (default)', '5-tuple', 'Source IP + Source Port + Dest IP + Dest Port + Protocol'],
    ['Client IP', '2-tuple', 'Source IP + Dest IP (same client always to same VM)'],
    ['Client IP and Protocol', '3-tuple', 'Source IP + Dest IP + Protocol'],
])

doc.add_heading('Load Balancing Alternatives', level=3)
add_table(doc, ['Service', 'Layer', 'Scope', 'Use Case'], [
    ['Azure Load Balancer', 'Layer 4', 'Regional', 'Non-HTTP traffic, TCP/UDP'],
    ['Application Gateway', 'Layer 7', 'Regional', 'HTTP/HTTPS, WAF, path-based routing'],
    ['Front Door', 'Layer 7', 'Global', 'Global HTTP/HTTPS, CDN, WAF'],
    ['Traffic Manager', 'DNS-based', 'Global', 'DNS routing, geographic distribution'],
])

# --- Module 18: Azure Application Gateway ---
doc.add_heading('Module 18: Azure Application Gateway', level=2)

add_bullet_list(doc, [
    'Operates at OSI Layer 7 (Application layer - HTTP/HTTPS)',
    'Web traffic load balancer with URL-based routing capabilities',
])

doc.add_heading('Components', level=3)
add_table(doc, ['Component', 'Description'], [
    ['Front-end IP', 'Public, private, or both IP addresses'],
    ['Listeners', 'Basic (single site) or Multi-site. Accept incoming connections.'],
    ['Routing rules', 'Bind listeners to back-end pools. Basic or path-based.'],
    ['Back-end pools', 'Web servers, VMSS, App Service, on-premises servers'],
    ['Health probes', 'HTTP 200 response = healthy'],
])

doc.add_heading('Key Features', level=3)
add_bullet_list(doc, [
    'Web Application Firewall (WAF): Protection against SQL injection, XSS, OWASP Core Rule Set',
    'Path-based routing: Route /images/* to one pool, /videos/* to another pool',
    'Multi-site routing: Host multiple websites on same Application Gateway',
    'TLS/SSL termination: Decrypt at gateway, send unencrypted to backend (offloads crypto work)',
    'Round-robin load balancing between backend servers',
    'Session stickiness (cookie-based affinity)',
    'Autoscaling based on traffic patterns',
    'Connection draining: Gracefully remove backend instances during updates',
    'WebSocket and HTTP/2 support',
    'URL redirect and URL rewrite capabilities',
])

# --- Module 19: Azure Network Watcher ---
doc.add_heading('Module 19: Azure Network Watcher', level=2)

add_bullet_list(doc, [
    'Automatically available when you create a VNet in a region',
    'IaaS monitoring only (not PaaS services)',
])

doc.add_heading('Network Watcher Tools', level=3)
add_table(doc, ['Category', 'Tool', 'Description'], [
    ['Monitoring', 'Topology', 'Visual diagram of VNet resources and relationships'],
    ['Monitoring', 'Connection Monitor', 'End-to-end connection monitoring between endpoints'],
    ['Diagnostic', 'IP Flow Verify', 'Check if a packet is allowed/denied and which NSG rule applies'],
    ['Diagnostic', 'NSG Diagnostics', 'Detailed NSG rule analysis for traffic flows'],
    ['Diagnostic', 'Next Hop', 'Determine next hop for a packet from a VM'],
    ['Diagnostic', 'Effective Security Rules', 'View all effective NSG rules applied to a NIC'],
    ['Diagnostic', 'Connection Troubleshoot', 'Check connectivity between source and destination'],
    ['Diagnostic', 'Packet Capture', 'Capture packets to/from a VM for analysis'],
    ['Diagnostic', 'VPN Troubleshoot', 'Diagnose VPN gateway and connection issues'],
    ['Traffic', 'Flow Logs', 'Log NSG traffic flow information'],
    ['Traffic', 'Traffic Analytics', 'Analyze flow logs for insights and patterns'],
])

# ============================================================
# PART 2 SECTION E: DEPLOY & MANAGE COMPUTE RESOURCES
# ============================================================
section_break(doc)
doc.add_heading('Section E: Deploy & Manage Compute Resources', level=2)

# --- Module 20: Introduction to Azure VMs ---
doc.add_heading('Module 20: Introduction to Azure Virtual Machines', level=2)

doc.add_heading('VM Planning Checklist', level=3)
add_bullet_list(doc, [
    'Network: Plan VNets, subnets, NSGs, IP addresses before creating VMs',
    'VM Name: Up to 64 chars (Linux) or 15 chars (Windows). Use consistent naming convention.',
    'Location: Choose region close to users. Affects pricing and available VM sizes.',
    'VM Size: Based on workload type (General purpose, Compute optimized, Memory optimized, etc.)',
    'Disks: OS disk + optional data disks. Separate data from OS for easier recovery.',
    'Operating System: Windows or Linux. Marketplace images or custom images available.',
])

doc.add_heading('VM Size Categories', level=3)
add_table(doc, ['Category', 'Description', 'Use Cases'], [
    ['General Purpose', 'Balanced CPU-to-memory', 'Dev/test, small-medium databases, low-medium traffic web servers'],
    ['Compute Optimized', 'High CPU-to-memory', 'Medium traffic web servers, network appliances, batch processing'],
    ['Memory Optimized', 'High memory-to-CPU', 'Relational databases, medium-large caches, in-memory analytics'],
    ['Storage Optimized', 'High disk throughput', 'Big data, SQL/NoSQL databases, data warehousing'],
    ['GPU', 'Heavy graphics/compute', 'Model training, deep learning, video editing'],
    ['High Performance Compute', 'Fastest CPUs', 'Molecular modeling, fluid dynamics, weather simulation'],
])

doc.add_heading('VM Disk Types', level=3)
add_table(doc, ['Disk Type', 'Max Size', 'Max IOPS', 'Max Throughput', 'Use Case'], [
    ['Ultra Disk', '65,536 GiB', '160,000', '4,000 MB/s', 'IO-intensive (SAP HANA, top-tier DBs)'],
    ['Premium SSD v2', '65,536 GiB', '80,000', '1,200 MB/s', 'Production workloads, low latency'],
    ['Premium SSD', '32,767 GiB', '20,000', '900 MB/s', 'Production, performance sensitive'],
    ['Standard SSD', '32,767 GiB', '6,000', '750 MB/s', 'Web servers, light enterprise apps, dev/test'],
    ['Standard HDD', '32,767 GiB', '2,000', '500 MB/s', 'Backup, noncritical, infrequent access'],
])

add_key_point(doc, 'Ultra Disk and Premium SSD v2 cannot be used as OS disks')

doc.add_heading('VM Pricing', level=3)
add_bullet_list(doc, [
    'Compute costs: Billed per-minute. Stop and deallocate to avoid charges.',
    'Storage costs: Charged separately even when VM is deallocated.',
    'Pay-as-you-go: Per second billing, no commitment, scale up/down anytime',
    'Reserved Instances (RI): 1 or 3 year commitment, up to 72% savings',
    'Azure Hybrid Benefit: Use existing Windows/Linux licenses to save on VM costs',
    'Linux VMs are cheaper (no OS license cost)',
])

doc.add_heading('VM Creation Options', level=3)
add_table(doc, ['Method', 'Description'], [
    ['Azure Portal', 'Web-based GUI, easiest for beginners'],
    ['Azure CLI', 'Cross-platform command-line (az vm create)'],
    ['Azure PowerShell', 'PowerShell cmdlets (New-AzVM)'],
    ['ARM/Bicep Templates', 'Infrastructure as Code, repeatable deployments'],
    ['Terraform', 'HCL syntax, multi-cloud IaC'],
    ['REST API / SDKs', 'Programmatic access (.NET, Java, Python, etc.)'],
])

doc.add_heading('VM Extensions', level=3)
add_bullet_list(doc, [
    'Small applications for post-deployment configuration and automation',
    'Install on VMs after initial deployment',
    'Examples: Custom Script Extension, DSC Extension, Diagnostics Extension',
])

doc.add_heading('VM Availability & Backup', level=3)
add_bullet_list(doc, [
    'Availability Zones: Protect against datacenter failure (3 zones per region)',
    'VM Scale Sets: Auto-scaling group of identical VMs',
    'Azure Load Balancer: Distribute traffic across VMs',
    'Azure Site Recovery: Replicate VMs to secondary region for DR',
    'Azure Backup: Backup and restore VMs with Recovery Services vault',
    'Auto-shutdown: Schedule VM shutdown to save costs',
])

# --- Module 21: Configure VM Availability ---
doc.add_heading('Module 21: Configure Virtual Machine Availability', level=2)

doc.add_heading('Maintenance & Downtime Types', level=3)
add_table(doc, ['Type', 'Description', 'Azure Response'], [
    ['Unplanned Hardware Maintenance', 'Azure predicts hardware failure', 'Live Migration to healthy host (brief pause)'],
    ['Unexpected Downtime', 'Hardware/infrastructure fails unexpectedly', 'Auto-heal: Migrate VM to healthy server (reboot)'],
    ['Planned Maintenance', 'Microsoft periodic platform updates', 'Updates rolled out by update domain (only one at a time)'],
])

doc.add_heading('Availability Sets', level=3)
add_bullet_list(doc, [
    'Logical grouping of VMs to prevent single point of failure',
    'All VMs should perform identical functionalities and have same software',
    'VM can only be added to availability set during creation',
    'To change availability set, must delete and recreate the VM',
])

doc.add_heading('Update Domains & Fault Domains', level=3)
add_table(doc, ['Feature', 'Update Domains', 'Fault Domains'], [
    ['Purpose', 'Group VMs rebooted together during updates', 'Group VMs sharing same physical hardware (power/network)'],
    ['Count', '1-20 (default: 5)', 'Up to 3'],
    ['Protection', 'Planned maintenance protection', 'Hardware failure protection'],
    ['Behavior', 'Only one UD rebooted at a time', 'Isolates against rack-level failures'],
])

add_key_point(doc, 'Update domain count is immutable after creation - must delete and recreate availability set to change')

doc.add_heading('Availability Zones', level=3)
add_bullet_list(doc, [
    'Physically separate locations within an Azure region',
    'Each zone has independent power, cooling, and networking',
    '3 availability zones per supported region',
    'Protects against entire datacenter failure',
    'Higher availability than availability sets',
])

doc.add_heading('Vertical vs Horizontal Scaling', level=3)
add_table(doc, ['Feature', 'Vertical (Scale Up/Down)', 'Horizontal (Scale Out/In)'], [
    ['What changes', 'VM size (CPU, memory, disk)', 'Number of VM instances'],
    ['Limitations', 'Hardware limits, requires reboot', 'Fewer limitations, thousands of instances possible'],
    ['Downtime', 'Requires VM stop/restart', 'No downtime'],
    ['Best for', 'Single VM needing more power', 'Variable workloads, high availability'],
])

doc.add_heading('VM Scale Sets', level=3)
add_bullet_list(doc, [
    'Deploy and manage a group of identical, auto-scaling VMs',
    'Supports Azure Load Balancer (L4) and Application Gateway (L7)',
    'Orchestration modes: Uniform (identical config) or Flexible (different sizes/images)',
    'Autoscale: Metric-based rules or schedule-based rules',
    'Spreading algorithm: Max spreading (recommended) vs Fixed spreading (exactly 5 FDs)',
])

doc.add_heading('Autoscale Configuration', level=3)
add_bullet_list(doc, [
    'Define minimum, maximum, and default instance count',
    'Scale-out rule: Increase instances when metric threshold exceeded (e.g., CPU > 70%)',
    'Scale-in rule: Decrease instances when metric drops (e.g., CPU < 30%)',
    'Always use BOTH scale-out and scale-in rules together',
    'Query duration: Look-back period for metric averaging',
    'Schedule-based scaling: Specific start/end dates and times',
])

# --- Module 22: Configure Azure App Service Plans ---
section_break(doc)
doc.add_heading('Module 22: Configure Azure App Service Plans', level=2)

doc.add_heading('App Service Plan Basics', level=3)
add_bullet_list(doc, [
    'Defines compute resources for web apps to run on',
    'Settings: Operating System (Windows/Linux), Region, Pricing Tier, VM Instances, VM Size',
    'Multiple apps can share the same App Service plan',
    'All apps in a plan run on the same VM instances',
    'Isolate apps into separate plans when: Resource-intensive, need independent scaling, need different region',
])

doc.add_heading('Pricing Tiers', level=3)
add_table(doc, ['Tier', 'Category', 'Features', 'Scale Out'], [
    ['Free F1', 'Shared compute', 'Dev/test only, no SLA, no custom domain', 'N/A'],
    ['Shared D1', 'Shared compute', 'Dev/test, custom domain, no SLA', 'N/A'],
    ['Basic B1-B3', 'Dedicated compute', 'Dev/test, manual scaling, no auto-scale', 'Up to 3 instances'],
    ['Standard S1-S3', 'Dedicated compute', 'Production, auto-scale, 5 staging slots, 10 daily backups', 'Up to 10 instances'],
    ['Premium P1V3-P3V3', 'Dedicated compute', 'Enhanced performance, 20 staging slots, 50 daily backups', 'Up to 30 instances'],
    ['Isolated I1V2-I3V2', 'Isolated', 'Network isolation, dedicated VNet, highest scale', 'Up to 200 instances'],
])

doc.add_heading('Scaling in App Service', level=3)
add_bullet_list(doc, [
    'Scale up: Change pricing tier (more CPU, memory, disk, features)',
    'Scale out: Increase number of VM instances running your app',
    'Autoscale: Automatic scale-out based on metric or schedule rules',
    'Automatic scaling (Elastic): PremiumV2/V3 only, HTTP traffic-based, platform-managed',
    'No redeployment needed when changing scale settings',
])

# --- Module 23: Configure Azure App Service ---
doc.add_heading('Module 23: Configure Azure App Service', level=2)

doc.add_heading('App Service Features', level=3)
add_bullet_list(doc, [
    'Multiple languages: ASP.NET, Java, Node.js, PHP, Python',
    'DevOps: CI/CD with Azure DevOps, GitHub, Bitbucket, Docker Hub',
    'Global scale with high availability',
    'Security and compliance: ISO, SOC, PCI compliant',
    'API and mobile features with CORS support',
    'Application templates from Azure Marketplace',
])

doc.add_heading('App Configuration Settings', level=3)
add_bullet_list(doc, [
    'Always On: Keep app loaded even without traffic (required for continuous WebJobs)',
    'Session affinity: Route client to same instance for session lifetime',
    'HTTPS Only: Redirect all HTTP to HTTPS',
])

doc.add_heading('Deployment Sources', level=3)
add_table(doc, ['Type', 'Sources'], [
    ['Automated (CI/CD)', 'GitHub (Actions), Bitbucket, Local Git, Azure Repos'],
    ['Manual', 'Remote Git, az webapp deploy CLI, Zip deploy'],
])

doc.add_heading('Deployment Slots', level=3)
add_bullet_list(doc, [
    'Available in Standard, Premium, and Isolated tiers',
    'Live apps with their own hostnames',
    'Swap content and configuration between slots (including production)',
    'Auto swap: Automatically swap to production after warm-up',
    'Validate changes in staging before swapping to production',
    'Easy rollback: Swap back immediately if issues found',
])

doc.add_heading('Swapped vs Slot-Specific Settings', level=3)
add_table(doc, ['Swapped (Follow Content)', 'Slot-Specific (Stay with Slot)'], [
    ['Language stack and version', 'Custom domain names'],
    ['App settings *', 'Non-public certificates, TLS/SSL settings'],
    ['Connection strings *', 'Scale settings'],
    ['Mounted storage accounts *', 'Always On setting'],
    ['Public certificates', 'IP restrictions'],
    ['WebJobs content', 'WebJobs schedulers'],
    ['Path mapping', 'Diagnostic settings'],
    ['', 'CORS'],
    ['', 'Virtual network integration'],
    ['', 'Managed identities'],
])

p = doc.add_paragraph()
p.add_run('* Can be configured to be slot-specific').font.size = Pt(9)

doc.add_heading('App Security', level=3)
add_bullet_list(doc, [
    'Built-in authentication: Entra ID, Google, Facebook, X (Twitter), Microsoft',
    'Runs in same environment as app code, but separately',
    'Options: Allow anonymous requests OR Allow only authenticated requests',
    'No SDK or code changes required',
])

doc.add_heading('Custom Domain Names', level=3)
add_bullet_list(doc, [
    'Default: <appname>.azurewebsites.net',
    'Steps: Reserve domain > Create DNS records (A or CNAME) > Enable in portal',
    'A record: Maps domain to IP address',
    'CNAME record: Maps domain to another domain (e.g., contoso.com > webapp.azurewebsites.net)',
    'Free managed TLS certificate available from App Service',
])

doc.add_heading('Backup & Restore', level=3)
add_bullet_list(doc, [
    'Available in Basic (production slot only), Standard, Premium, Isolated tiers',
    'Requires Azure storage account in same subscription',
    'Backs up: App config, file content, connected databases',
    'Each backup: Zip file + XML manifest',
    'Full or partial backups supported',
    'Max backup size: 10 GB (app + database)',
    'Manual or scheduled backups',
])

doc.add_heading('Azure Application Insights', level=3)
add_bullet_list(doc, [
    'Feature of Azure Monitor for monitoring live applications',
    'Supports .NET, Node.js, Java EE, and more',
    'Monitors: Request rates, response times, failure rates, dependency performance',
    'Tracks: Page views, user/session counts, performance counters, custom events',
    'Diagnostic trace logs for correlating trace events with requests',
])

# --- Module 24: Configure Azure Container Instances ---
section_break(doc)
doc.add_heading('Module 24: Configure Azure Container Instances', level=2)

doc.add_heading('Containers vs Virtual Machines', level=3)
add_table(doc, ['Feature', 'Containers', 'Virtual Machines'], [
    ['Isolation', 'Lightweight (process-level)', 'Complete (hardware-level, strongest security)'],
    ['OS', 'Shares host OS kernel, user-mode only', 'Full OS including kernel'],
    ['Startup', 'Seconds', 'Minutes'],
    ['Resource usage', 'Minimal (only needed services)', 'Full OS overhead'],
    ['Deployment', 'Docker CLI, orchestrators (AKS)', 'Portal, PowerShell, Hyper-V'],
    ['Storage', 'Azure Disks, Azure Files', 'VHDs, SMB shares'],
    ['Fault tolerance', 'Rapid recreation on another node', 'Failover with OS restart'],
])

doc.add_heading('Azure Container Instances (ACI)', level=3)
add_bullet_list(doc, [
    'Fastest and simplest way to run containers in Azure',
    'No VM management required (serverless)',
    'Public IP connectivity and DNS names for internet access',
    'Custom container sizes (CPU cores and memory)',
    'Supports Linux and Windows containers',
    'Persistent storage with Azure Files volumes',
])

doc.add_heading('Container Groups', level=3)
add_bullet_list(doc, [
    'Top-level resource in ACI - collection of containers on same host machine',
    'Similar to a Kubernetes pod',
    'Share lifecycle, resources, local network, and storage volumes',
    'Deployment via: ARM templates, Bicep, YAML files',
    'Shared external-facing IP address',
    'Port mapping NOT supported (containers share port namespace)',
    'When deleted: IP address and FQDN are released',
])

doc.add_heading('Multi-Container Group Use Cases', level=3)
add_bullet_list(doc, [
    'Web app updates: One container serves app, another pulls latest content',
    'Log data collection: App container outputs logs, logging container writes to storage',
    'App monitoring: Monitoring container checks health of application container',
    'Front-end + Back-end: Web container + API/data service container',
])

doc.add_heading('Container Service Comparison', level=3)
add_table(doc, ['Service', 'Best For', 'Key Feature'], [
    ['Azure Container Instances (ACI)', 'Isolated, short-lived tasks', 'Simplest, no orchestration'],
    ['Azure Container Apps (ACA)', 'Serverless microservices', 'Event-driven scaling, Dapr, KEDA'],
    ['Azure Kubernetes Service (AKS)', 'Complex orchestration', 'Full Kubernetes control, enterprise workloads'],
])

# ============================================================
# PART 2 SECTION F: MONITOR & BACK UP AZURE RESOURCES
# ============================================================
section_break(doc)
doc.add_heading('Section F: Monitor & Back up Azure Resources', level=2)

# --- Module 25: Introduction to Azure Backup ---
doc.add_heading('Module 25: Introduction to Azure Backup', level=2)

doc.add_heading('Azure Backup Overview', level=3)
add_bullet_list(doc, [
    'Zero-infrastructure backup: No backup servers to deploy or manage',
    'Backup Center: Single unified console for managing all backups at scale',
    'Automatic protection against ransomware, malicious admins, accidental deletions',
])

doc.add_heading('Key Concepts', level=3)
add_table(doc, ['Concept', 'Description'], [
    ['RTO (Recovery Time Objective)', 'Maximum acceptable time to restore after disruption'],
    ['RPO (Recovery Point Objective)', 'Maximum acceptable data loss measured in time'],
    ['Backup Center', 'Single pane of glass for managing, monitoring, governing, and optimizing backups'],
])

doc.add_heading('Supported Workloads', level=3)
add_bullet_list(doc, [
    'Azure VMs (Windows and Linux)',
    'Azure Managed Disks',
    'Azure Files shares',
    'SQL Server on Azure VMs',
    'SAP HANA on Azure VMs',
    'Azure Database for PostgreSQL',
    'Azure Blobs',
    'Azure Database for MySQL',
    'Azure Kubernetes Service (AKS)',
    'On-premises: Files, folders, system state via MARS agent',
])

doc.add_heading('Backup Types', level=3)
add_table(doc, ['Type', 'Description'], [
    ['Full backup', 'Complete copy of all data'],
    ['Differential backup', 'Only data changed since last full backup'],
    ['Transaction log backup', 'Database transaction logs for point-in-time recovery'],
    ['Selective disk backup', 'Back up only specific disks attached to VM'],
])

doc.add_heading('Backup Access Tiers', level=3)
add_table(doc, ['Tier', 'Speed', 'Cost', 'Description'], [
    ['Snapshot Tier', 'Fastest', 'Higher', 'Stored locally for instant restore. Max 5 days retention.'],
    ['Vault-Standard Tier', 'Medium', 'Medium', 'Transferred to vault for longer retention.'],
    ['Archive Tier', 'Slowest', 'Cheapest', 'Long-term retention with lowest cost.'],
])

doc.add_heading('Vault Types', level=3)
add_table(doc, ['Vault', 'Supported Workloads'], [
    ['Recovery Services Vault', 'Azure VMs, SQL on VM, Azure Files, on-premises (MARS/MABS/DPM)'],
    ['Backup Vault', 'Azure Disks, Azure Blobs, Azure Database for PostgreSQL, AKS'],
])

doc.add_heading('Security Features', level=3)
add_bullet_list(doc, [
    'Soft delete: Retains backup data 14 days after deletion (free)',
    'Enhanced soft delete: Longer retention period, always-on option',
    'RBAC: Segregate duties and limit access',
    'Encryption: Microsoft-managed or customer-managed keys',
    'Storage redundancy: LRS, GRS, ZRS for vault storage',
    'Backup Explorer: Azure Monitor workbook for monitoring backup estate',
])

# --- Module 26: Protect VMs with Azure Backup ---
doc.add_heading('Module 26: Protect Virtual Machines with Azure Backup', level=2)

doc.add_heading('Azure Backup vs Azure Site Recovery', level=3)
add_table(doc, ['Feature', 'Azure Backup', 'Azure Site Recovery'], [
    ['Purpose', 'Maintain copies of stateful data to go back in time', 'Real-time replication for failover'],
    ['Use Case', 'Accidental data loss, corruption, ransomware', 'Network/power outages, regional disasters'],
    ['Approach', 'Point-in-time snapshots', 'Near real-time replication'],
])

doc.add_heading('Backup Features', level=3)
add_bullet_list(doc, [
    'Zero-infrastructure: No backup servers to manage',
    'Long-term retention: Meet compliance and audit requirements',
    'RBAC: Segregate duties for backup management',
    'Encryption: Microsoft-managed or customer-managed keys (Key Vault)',
    'No internet connectivity required for Azure VMs (Azure backbone only)',
    'Soft delete: 14 extra days retention after deletion',
    'Enhanced soft delete: Customizable retention, always-on protection',
    'High availability: LRS (noncritical), GRS (recommended), ZRS (high availability)',
])

doc.add_heading('Recovery Services Vault', level=3)
add_bullet_list(doc, [
    'Storage management entity for backup data',
    'Handles backup and restore operations',
    'RBAC boundary for secure access',
    'Data transferred to Azure Backup storage in separate fault domain',
])

doc.add_heading('VM Backup Process', level=3)
add_bullet_list(doc, [
    '1. Backup job starts per policy schedule',
    '2. Backup extension installed on VM (VMSnapshot for Windows, VMSnapshotLinux for Linux)',
    '3. Snapshot taken of all disks in parallel',
    '4. Data transferred to vault (only delta/changed blocks after first backup)',
    '5. Snapshot data may take hours to transfer at peak times',
    '6. Total backup time < 24 hours for daily policies',
])

doc.add_heading('Snapshot Consistency Levels', level=3)
add_table(doc, ['Level', 'Description', 'How'], [
    ['Application consistent', 'Captures VM + memory + pending I/O', 'VSS (Windows) or custom scripts (Linux)'],
    ['File system consistent', 'Captures disk state without memory', 'When VSS fails or scripts fail'],
    ['Crash consistent', 'Captures disk only (no I/O or memory)', 'When VM is shut down during backup'],
])

doc.add_heading('Backup Policy', level=3)
add_bullet_list(doc, [
    'Frequency: Daily or weekly (Enhanced policy supports hourly)',
    'Retention: Days, weeks, months, or years',
    'Two access tiers: Snapshot tier (instant, max 5 days) and Vault tier (longer retention)',
    'Selective disk backup: Back up only specific disks (Enhanced policy)',
])

doc.add_heading('Restore Options', level=3)
add_table(doc, ['Restore Type', 'Description'], [
    ['Create a new VM', 'Quick VM creation from restore point. Must be same region as source.'],
    ['Restore disk', 'Restore disk, then create/customize VM. Provides ARM template.'],
    ['Replace existing', 'Replace disk on existing VM. VM must still exist. Takes snapshot before replacing.'],
    ['Cross Region', 'Restore to secondary (paired) region. Create VM or Restore Disk only.'],
    ['Cross Subscription', 'Restore to different subscription in same tenant. Managed VMs only.'],
    ['Cross Zonal', 'Restore to different availability zone. Managed VMs only. Requires ZRS vault.'],
    ['Selective Disk', 'Restore specific disks from recovery point.'],
])

doc.add_heading('File Recovery', level=3)
add_bullet_list(doc, [
    'Mount snapshot on target machine using iSCSI initiator',
    'Recover individual files without restoring entire VM',
])

doc.add_heading('Encrypted VM Restore Limitations', level=3)
add_bullet_list(doc, [
    'Only standalone key encryption supported (not keys part of certificate)',
    'File/folder level restore not supported (must restore entire VM)',
    'Replace existing VM option not available for encrypted VMs',
])

# --- Module 27: Monitor Azure VMs with Azure Monitor ---
doc.add_heading('Module 27: Monitor Azure VMs with Azure Monitor', level=2)

doc.add_heading('Azure Monitor Overview', level=3)
add_bullet_list(doc, [
    'Comprehensive monitoring solution for Azure and non-Azure resources',
    'Two main features: Azure Monitor Metrics and Azure Monitor Logs',
    'Metrics: Numerical values at predetermined intervals, retained 93 days',
    'Logs: Recorded system events with timestamps, stored in Log Analytics workspace',
])

doc.add_heading('VM Monitoring Layers', level=3)
add_table(doc, ['Layer', 'What to Monitor', 'How'], [
    ['Host VM', 'CPU, disk, network, availability', 'Automatic - built-in metrics and Activity logs'],
    ['Guest OS', 'OS performance, events', 'Azure Monitor Agent + DCR'],
    ['Client Workloads', 'Application performance', 'VM insights + custom DCRs'],
    ['Applications', 'App behavior and usage', 'Application Insights'],
])

doc.add_heading('VM Host Monitoring (Automatic)', level=3)
add_bullet_list(doc, [
    'Built-in metrics: VM availability, CPU usage, OS disk usage, Network operations, Disk operations/sec',
    'Recommended alert rules: Predefined rules for CPU, memory, disk, network, VM availability',
    'Activity logs: VM startup, modifications, RBAC changes (automatically recorded)',
    'Boot diagnostics: Screenshots and serial log from VM boot sequence',
])

doc.add_heading('Metrics Explorer', level=3)
add_bullet_list(doc, [
    'UI for exploring and analyzing VM metrics',
    'Plot multiple metrics on one graph',
    'Aggregation functions: Count, Average, Maximum, Minimum, Sum',
    'Flexible time ranges: 30 minutes to 30 days, custom ranges',
    'Split by dimension to show per-VM data',
    'Pin charts to dashboards or send to workbooks',
])

doc.add_heading('VM Insights', level=3)
add_bullet_list(doc, [
    'Simplified VM client monitoring onboarding',
    'Automatically installs Azure Monitor Agent',
    'Creates preconfigured DCR for common performance counters',
    'Predefined performance charts and workbooks',
    'Optional: Process and dependency mapping (Map feature)',
    'Data sent to Log Analytics workspace (view via Insights, not Metrics Explorer)',
])

doc.add_heading('Data Collection Rules (DCRs)', level=3)
add_bullet_list(doc, [
    'Define what data to collect and where to send it',
    'VM insights creates one automatically for performance counters',
    'Create custom DCRs for event logs and additional performance counters',
    'Associate a DCR with one or multiple VMs',
    'Data collection endpoint required for log data',
    'Supports: Linux Syslog, Windows Event Logs, custom performance counters',
])

doc.add_heading('Log Analytics & KQL', level=3)
add_bullet_list(doc, [
    'Log data stored in Log Analytics workspace',
    'Query using Kusto Query Language (KQL)',
    'Access from: VM > Logs, or Azure Monitor > Logs',
    'Example query: Syslog | where SeverityLevel == "warning"',
    'Scope queries to specific VM or workspace',
])

# ============================================================
# FINAL SECTION: QUICK REFERENCE CARDS
# ============================================================
section_break(doc)
doc.add_heading('Quick Reference: Key Comparisons', level=1)

doc.add_heading('RBAC vs Azure Policy', level=3)
add_table(doc, ['Feature', 'RBAC', 'Azure Policy'], [
    ['Focus', 'User actions/permissions', 'Resource properties/compliance'],
    ['Scope', 'What users can DO', 'What resources LOOK LIKE'],
    ['Default', 'Deny all (allow model)', 'Allow all (deny/audit model)'],
    ['Example', 'Allow user to create VMs', 'All VMs must use managed disks'],
])

doc.add_heading('Service Endpoints vs Private Endpoints', level=3)
add_table(doc, ['Feature', 'Service Endpoints', 'Private Endpoints'], [
    ['IP Address', 'Public IP (Azure backbone)', 'Private IP from your VNet'],
    ['Traffic path', 'Azure backbone network', 'Private link'],
    ['Access from on-premises', 'Not natively (needs workaround)', 'Yes, via VPN/ExpressRoute'],
    ['Cost', 'Free', 'Charged per hour + per GB processed'],
    ['Security', 'Good', 'Best (full network isolation)'],
])

doc.add_heading('Backup Vault vs Recovery Services Vault', level=3)
add_table(doc, ['Feature', 'Recovery Services Vault', 'Backup Vault'], [
    ['Azure VMs', 'Yes', 'No'],
    ['SQL/SAP HANA on VM', 'Yes', 'No'],
    ['Azure Files', 'Yes', 'No'],
    ['On-premises (MARS/MABS)', 'Yes', 'No'],
    ['Azure Disks', 'No', 'Yes'],
    ['Azure Blobs', 'No', 'Yes'],
    ['Azure PostgreSQL', 'No', 'Yes'],
    ['AKS', 'No', 'Yes'],
])

doc.add_heading('Load Balancer vs Application Gateway', level=3)
add_table(doc, ['Feature', 'Azure Load Balancer', 'Application Gateway'], [
    ['OSI Layer', 'Layer 4 (Transport)', 'Layer 7 (Application)'],
    ['Protocol', 'TCP/UDP', 'HTTP/HTTPS'],
    ['Routing', 'IP + Port hash', 'URL path, host header, cookies'],
    ['WAF', 'No', 'Yes (OWASP CRS)'],
    ['SSL Termination', 'No', 'Yes'],
    ['Scope', 'Regional', 'Regional'],
])

# ============================================================
# COMPREHENSIVE ASSESSMENT ANSWERS - ALL MODULES
# ============================================================
section_break(doc)
doc.add_heading('ALL MODULE ASSESSMENTS & KNOWLEDGE CHECKS', level=1)
p = doc.add_paragraph()
run = p.add_run('This section contains every assessment and knowledge check answer from all modules covered in both AZ-305 and AZ-104 learning paths.')
run.font.size = Pt(11)
run.italic = True
doc.add_paragraph()

# --- AZ-305 Design Governance Assessment ---
doc.add_heading('AZ-305: Design Governance - Assessment', level=2)
add_assessment_box(doc, [
    ('Which Azure AD feature supports requesting access to groups or application access packages?',
     'Entitlement management'),
    ('Which Azure RBAC role allows managing user access but not resources?',
     'User Access Administrator'),
    ('What is the maximum number of levels of depth for management groups (not including root or subscription)?',
     '6 levels of depth'),
    ('Are VNets shared across subscriptions?',
     'No - Virtual networks cannot be shared across subscriptions'),
    ('Can resource groups be nested?',
     'No - Resource groups cannot be nested inside other resource groups'),
    ('Are resource tags inherited from resource group to resources?',
     'No - Tags are NOT inherited. Use Azure Policy to enforce tag inheritance.'),
    ('What is the recommended management group hierarchy depth?',
     'Flat hierarchy of 3-4 levels'),
])

# --- AZ-305 Design Authentication & Authorization Assessment ---
doc.add_heading('AZ-305: Design Authentication & Authorization - Assessment', level=2)
add_assessment_box(doc, [
    ('What provides single set of credentials across cloud and on-premises?',
     'Hybrid identity with Microsoft Entra Connect'),
    ('What are the 4 pillars of IAM?',
     'Unified identity management, Seamless user experience, Secure adaptive access, Simplified identity governance'),
    ('What license is required for Conditional Access?',
     'Microsoft Entra ID P1 or P2'),
    ('What is the difference between system-assigned and user-assigned managed identities?',
     'System-assigned: tied to resource lifecycle, cannot be shared. User-assigned: standalone, can be shared across resources.'),
    ('What are the 3 types of service principals?',
     'Application (most common), Managed Identity, Legacy'),
    ('What are the two Key Vault tiers?',
     'Standard (software encryption) and Premium (HSM-backed keys)'),
    ('What is the relationship between app objects and service principals?',
     'App object is 1:many with service principals - one app object, service principal in each tenant where app is used'),
    ('What are the 3 reviewer types for Access Reviews?',
     'Resource owners, Delegates (managers), End users (self-review)'),
    ('What are user risk vs sign-in risk examples?',
     'User risk: Leaked credentials, Threat intelligence. Sign-in risk: Anonymous IP, Atypical travel, Password spray, Malicious IP, Anomalous token, Verified threat actor IP.'),
    ('Is Azure AD B2C available for new customers?',
     'No - Azure AD B2C no longer available for new customers since May 2025. Use Microsoft Entra External ID instead. Existing B2C supported until May 2030.'),
])

# --- AZ-104 Microsoft Entra ID Assessment ---
doc.add_heading('AZ-104: Microsoft Entra ID - Assessment', level=2)
add_assessment_box(doc, [
    ('What type of service is Microsoft Entra ID?',
     'PaaS (Platform as a Service) - Identity as a Service'),
    ('What query protocol does Entra ID use instead of LDAP?',
     'REST API over HTTP/HTTPS'),
    ('What authentication protocols does Entra ID use?',
     'SAML, OpenID Connect, OAuth (NOT Kerberos/NTLM like traditional AD)'),
    ('What does Entra ID P2 add over P1?',
     'Identity Protection (risk-based policies), Privileged Identity Management (PIM), Access Reviews, Entitlement Management'),
    ('What is Microsoft Entra Domain Services?',
     'Managed domain services (domain join, Group Policy, LDAP, Kerberos/NTLM) without deploying domain controllers'),
])

# --- AZ-104 Create & Manage Identities Assessment ---
doc.add_heading('AZ-104: Create & Manage Identities - Assessment', level=2)
add_assessment_box(doc, [
    ('What are the three user account types?',
     'Cloud Identity (Entra ID only), Directory-synced Identity (from on-prem AD), Guest User (B2B)'),
    ('What is the difference between Security and Microsoft 365 groups?',
     'Security: manage access, can contain users/devices/service principals/groups. M365: collaboration, shared mailbox/calendar/files, members can only be users.'),
    ('What license is required for Dynamic groups?',
     'Microsoft Entra ID P1'),
    ('What are the three device identity options?',
     'Entra Registered (BYOD, personal), Entra Joined (organization-owned, cloud), Hybrid Entra Joined (both on-prem AD and Entra ID)'),
])

# --- AZ-104 Azure Policy Assessment ---
doc.add_heading('AZ-104: Azure Policy - Assessment', level=2)
add_assessment_box(doc, [
    ('What are the 6 Azure Policy resources?',
     'Policy Definition, Policy Initiative, Policy Assignment, Policy Exemption, Policy Attestation, Policy Remediation'),
    ('How often does Azure Policy evaluate resources automatically?',
     'Every 24 hours'),
    ('What is the difference between Deny and Audit effects?',
     'Deny: Prevents resource creation/update BEFORE deployment. Audit: Creates warning in Activity Log AFTER deployment.'),
    ('What is enforcementMode?',
     'Default (enforce) or DoNotEnforce (audit only without blocking). Used for safe deployment testing.'),
    ('What is a Policy Initiative?',
     'A collection of policy definitions grouped together for a common goal (e.g., "Security Center" initiative)'),
])

# --- AZ-104 Azure RBAC Assessment ---
doc.add_heading('AZ-104: Azure RBAC - Assessment', level=2)
add_assessment_box(doc, [
    ('What are the 3 components of a role assignment?',
     'Security Principal (who) + Role Definition (what) + Scope (where)'),
    ('What is the RBAC access model?',
     'Allow model - permissions are additive. Effective permissions = sum of all role assignments.'),
    ('What are the 4 key built-in roles?',
     'Owner (full + manage access), Contributor (full but no access management), Reader (view only), User Access Administrator (manage access only)'),
    ('What is the scope hierarchy?',
     'Management group > Subscription > Resource group > Resource. Permissions inherit downward.'),
])

# --- AZ-104 Microsoft Entra SSPR Assessment ---
doc.add_heading('AZ-104: Microsoft Entra SSPR - Knowledge Check', level=2)
add_assessment_box(doc, [
    ('Which license is required for SSPR?',
     'Microsoft Entra ID P1 or P2 (paid subscription required)'),
    ('How many authentication methods do admins need for SSPR?',
     'Always 2 methods - this cannot be changed. Regular users: 1 or 2 (admin configurable).'),
    ('Can admins use security questions for SSPR?',
     'No - Security questions are available for users ONLY, not for admins, and not for MFA'),
    ('What is the SSPR registration URL?',
     'https://aka.ms/ssprsetup (combined SSPR+MFA: https://mysignins.microsoft.com)'),
    ('What are the SSPR flow steps?',
     'Localization → CAPTCHA verification → Authentication (user ID) → Password reset → Notification'),
    ('What is password writeback?',
     'Syncs cloud password changes back to on-premises AD via Entra Connect or Cloud Sync. Requires P1.'),
    ('What are the custom branding image requirements?',
     'Background: PNG/JPG, 1920x1080 px, max 300 KB. Logo: PNG/JPG, 32x32 px, max 5 KB.'),
])

# --- AZ-104 Configure Storage Accounts Assessment ---
doc.add_heading('AZ-104: Configure Storage Accounts - Assessment', level=2)
add_assessment_box(doc, [
    ('Which replication strategy provides the highest durability and availability?',
     'RA-GZRS (Read-Access Geo-Zone Redundant Storage) - 16 nines durability + read access to secondary region'),
    ('What is the default endpoint format for blob storage?',
     'https://<storage-account>.blob.core.windows.net'),
    ('Which storage account type is the default and recommended for most scenarios?',
     'Standard general-purpose v2 (GPv2)'),
    ('What are the two secure endpoint options for storage accounts?',
     'Service Endpoints (free, Azure backbone) and Private Endpoints (private IP from VNet)'),
    ('How many copies does LRS maintain?',
     '3 copies within a single datacenter'),
    ('Can you configure a custom domain for a storage account?',
     'Yes - via CNAME mapping (indirect) or direct mapping with asverify subdomain'),
])

# --- AZ-104 Configure Azure Blob Storage Assessment ---
doc.add_heading('AZ-104: Configure Azure Blob Storage - Assessment', level=2)
add_assessment_box(doc, [
    ('What is the minimum retention period for Cool tier?',
     '30 days - early deletion fee applies if deleted before 30 days'),
    ('What is the minimum retention period for Archive tier?',
     '180 days - data must be rehydrated before accessing (offline tier)'),
    ('Which blob type is used for VM disks (VHDs)?',
     'Page Blob - supports random read/write operations, up to 8 TB'),
    ('What is required for blob object replication?',
     'Blob versioning must be enabled on BOTH source and destination accounts. Snapshots are NOT supported.'),
    ('What are the three container access levels?',
     'Private (no anonymous access), Blob (anonymous read for blobs only), Container (anonymous read for blobs + container listing)'),
    ('How do storage costs change across tiers?',
     'Storage cost DECREASES from Hot to Archive, but Access cost INCREASES from Hot to Archive'),
    ('What tools are available for managing blob storage?',
     'Azure Storage Explorer (GUI), AzCopy (CLI), Data Box Disk (physical transfer)'),
])

# --- AZ-104 Configure Azure Storage Security Assessment ---
doc.add_heading('AZ-104: Configure Azure Storage Security - Assessment', level=2)
add_assessment_box(doc, [
    ('What encryption standard does Azure Storage use?',
     '256-bit AES encryption - automatic, enabled by default, cannot be disabled'),
    ('What is the most secure type of SAS?',
     'User Delegation SAS - secured with Entra ID credentials + SAS token. Blob storage only.'),
    ('What are the three types of SAS?',
     'User Delegation SAS, Service SAS, Account SAS'),
    ('What are the four authorization strategies for Azure Storage?',
     'Microsoft Entra ID (recommended), Shared Key, Shared Access Signature (SAS), Anonymous'),
    ('What is a stored access policy?',
     'Server-side policy for a Service SAS that allows you to revoke or modify SAS permissions without regenerating storage keys'),
    ('What are the components of Defender for Storage?',
     'Malware scanning, Sensitive data threat detection, Activity-based threat detection (anomalous access patterns)'),
    ('What is infrastructure encryption (double encryption)?',
     'Additional encryption layer at storage service level using a different algorithm/key, on top of default SSE'),
])

# --- AZ-104 Configure Azure Files Assessment ---
doc.add_heading('AZ-104: Configure Azure Files - Assessment', level=2)
add_assessment_box(doc, [
    ('What protocols does Azure Files support?',
     'SMB (port 445), NFS, and HTTP (REST API)'),
    ('What are the storage tiers for file shares?',
     'Premium (SSD/FileStorage), Transaction Optimized (HDD/GPv2), Hot (HDD/GPv2), Cool (HDD/GPv2)'),
    ('What are the three authentication methods for Azure Files?',
     'Identity-based (AD DS, Entra Domain Services, Entra Kerberos), Storage account access key, Shared Access Signature (SAS)'),
    ('What are file share snapshot characteristics?',
     'Incremental (only changed data saved), Read-only, Up to 200 per share, Captured at share level, Retrieved at file level'),
    ('What is the soft delete retention range?',
     '1 to 365 days. Enabled at storage account level (applies to all file shares).'),
    ('What are the 5 components of Azure File Sync?',
     'Storage Sync Service, Sync group, Cloud endpoint, Server endpoint, Azure File Sync Agent'),
    ('How many sync groups per Storage Sync Service?',
     'Up to 100 sync groups per Storage Sync Service, up to 100 server endpoints per sync group'),
])

# --- AZ-104 Configure Virtual Networks Assessment ---
doc.add_heading('AZ-104: Configure Virtual Networks - Assessment', level=2)
add_assessment_box(doc, [
    ('How many IP addresses are reserved per subnet in Azure?',
     '5 reserved IPs: x.x.x.0 (network), x.x.x.1 (gateway), x.x.x.2-3 (Azure DNS), x.x.x.255 (broadcast)'),
    ('What is the usable IP count for a /29 subnet?',
     '3 usable IPs (8 total - 5 reserved = 3)'),
    ('What is the allocation type for Standard SKU public IP?',
     'Always static. Zone-redundant by default. Closed to inbound traffic by default (must use NSG to allow).'),
    ('What are the private IP address ranges?',
     '10.0.0.0/8, 172.16.0.0/12, 192.168.0.0/16'),
    ('Can VNet address spaces overlap when peering?',
     'No - Address spaces must NOT overlap for VNet peering to work'),
])

# --- AZ-104 Configure NSGs Assessment ---
doc.add_heading('AZ-104: Configure NSGs - Assessment', level=2)
add_assessment_box(doc, [
    ('What is the priority range for NSG security rules?',
     '100 to 4096. Lower number = higher priority.'),
    ('What is the inbound NSG processing order?',
     'Subnet NSG evaluated FIRST, then NIC NSG. Traffic must be allowed by BOTH.'),
    ('What is the outbound NSG processing order?',
     'NIC NSG evaluated FIRST, then Subnet NSG.'),
    ('What are the default NSG inbound rules?',
     'AllowVNetInBound (65000), AllowAzureLoadBalancerInBound (65001), DenyAllInBound (65500)'),
    ('What are the default NSG outbound rules?',
     'AllowVNetOutBound (65000), AllowInternetOutBound (65001), DenyAllOutBound (65500)'),
    ('What are Application Security Groups (ASGs)?',
     'Group VMs by workload (e.g., WebServers, DbServers). Use as source/destination in NSG rules instead of explicit IPs.'),
    ('What are service tags?',
     'Predefined labels for Azure services (Internet, VirtualNetwork, AzureLoadBalancer, Storage, Sql). Simplify NSG rules.'),
])

# --- AZ-104 Azure DNS Knowledge Check ---
doc.add_heading('AZ-104: Host Domain on Azure DNS - Knowledge Check', level=2)
add_assessment_box(doc, [
    ('What is domain delegation?',
     'Update NS records at your domain registrar to point to Azure DNS name servers'),
    ('Can CNAME records be used at zone apex?',
     'No - CNAME cannot be used at zone apex. Use Alias records instead.'),
    ('What are alias records?',
     'Point directly to Azure resources (Traffic Manager, CDN, Public IP, Front Door). Auto-update when IP changes. Prevent dangling DNS.'),
    ('What is a private DNS zone used for?',
     'Name resolution within virtual networks without custom DNS solution. Supports auto-registration of VM DNS records.'),
    ('Does Azure DNS support domain name purchasing?',
     'No - Azure DNS only hosts DNS zones. Purchase domains via App Service domains or third-party registrars.'),
])

# --- AZ-104 VNet Peering Assessment ---
doc.add_heading('AZ-104: Configure VNet Peering - Assessment', level=2)
add_assessment_box(doc, [
    ('Is VNet peering transitive?',
     'No - Peering is nontransitive. VNet A <> B and B <> C does NOT mean A <> C. Use UDR/service chaining to extend.'),
    ('What are the two types of VNet peering?',
     'Regional peering (same region) and Global peering (different regions)'),
    ('What is gateway transit?',
     'Share VPN/ExpressRoute gateway through peered VNet. Hub VNet has gateway, spoke VNets use "Use remote gateways" setting.'),
    ('What is the peering status flow?',
     'Initiated > Connected. Both sides must complete peering for traffic to flow.'),
    ('What is hub-and-spoke topology?',
     'Central hub VNet with shared services (VPN gateway, firewall), spoke VNets for individual workloads. Connected via peering.'),
])

# --- AZ-104 Traffic Flow Knowledge Check ---
doc.add_heading('AZ-104: Manage Traffic Flow - Knowledge Check', level=2)
add_assessment_box(doc, [
    ('What is the route priority order?',
     '1st: Longest prefix match. 2nd: UDR > BGP > System routes.'),
    ('What is an NVA?',
     'Network Virtual Appliance - VM performing network functions: Firewall, WAN optimizer, Router, IDS/IPS. IP forwarding must be enabled.'),
    ('What is the "None" next hop type?',
     'Drop/blackhole traffic - packets are discarded'),
    ('What are system routes?',
     'Automatically created by Azure for VNet communication, Internet access, and blocking certain traffic. Cannot be deleted.'),
])

# --- AZ-104 Azure Load Balancer Assessment ---
doc.add_heading('AZ-104: Azure Load Balancer - Assessment', level=2)
add_assessment_box(doc, [
    ('At which OSI layer does Azure Load Balancer operate?',
     'Layer 4 (Transport layer - TCP/UDP)'),
    ('What is the default session persistence (distribution mode)?',
     'None - 5-tuple hash (Source IP + Source Port + Dest IP + Dest Port + Protocol)'),
    ('What are health probe types?',
     'TCP, HTTP, HTTPS'),
    ('What are HA (High Availability) ports?',
     'Protocol=All, Port=0. Load balance ALL TCP/UDP flows on ALL ports. Used with Internal Load Balancer.'),
    ('What are the load balancing alternatives?',
     'Load Balancer (L4 Regional), Application Gateway (L7 Regional), Front Door (L7 Global), Traffic Manager (DNS-based Global)'),
])

# --- AZ-104 Application Gateway Assessment ---
doc.add_heading('AZ-104: Azure Application Gateway - Assessment', level=2)
add_assessment_box(doc, [
    ('At which OSI layer does Application Gateway operate?',
     'Layer 7 (Application layer - HTTP/HTTPS)'),
    ('What is WAF?',
     'Web Application Firewall - Protection against SQL injection, XSS, and other OWASP Core Rule Set vulnerabilities'),
    ('What are the two routing types?',
     'Path-based routing (/images/* to pool A, /videos/* to pool B) and Multi-site routing (multiple websites on same gateway)'),
    ('What listener types are available?',
     'Basic (single site) and Multi-site (multiple hostnames on same listener)'),
    ('What HTTP response indicates healthy in health probes?',
     'HTTP 200 OK'),
    ('What is TLS/SSL termination?',
     'Decrypt HTTPS at gateway, send unencrypted HTTP to backend servers. Offloads cryptographic processing.'),
    ('What is connection draining?',
     'Gracefully remove backend instances during planned updates. Existing connections complete, no new connections sent.'),
])

# --- AZ-104 Azure Network Watcher Assessment ---
doc.add_heading('AZ-104: Azure Network Watcher - Assessment', level=2)
add_assessment_box(doc, [
    ('What are the 3 categories of Network Watcher tools?',
     'Monitoring (Topology, Connection Monitor), Diagnostic (IP Flow Verify, NSG Diagnostics, Next Hop, Effective Security Rules, Connection Troubleshoot, Packet Capture, VPN Troubleshoot), Traffic (Flow Logs, Traffic Analytics)'),
    ('When is Network Watcher automatically available?',
     'Automatically created when you create a VNet in a region'),
    ('Does Network Watcher support PaaS services?',
     'No - IaaS monitoring only (VMs, VNets, NSGs). Not for PaaS services.'),
    ('What does IP Flow Verify do?',
     'Checks if a packet is allowed or denied to/from a VM and identifies which NSG rule is responsible'),
    ('What does Next Hop do?',
     'Determines the next hop for a packet from a VM - helps troubleshoot routing issues'),
])

# --- AZ-104 Introduction to Azure Backup Assessment ---
doc.add_heading('AZ-104: Introduction to Azure Backup - Assessment', level=2)
add_assessment_box(doc, [
    ('Which backup tier provides the fastest recovery?',
     'Snapshot Tier - stored locally for instant restore, maximum 5 days retention'),
    ('What provides single unified management experience for backups?',
     'Backup Center - discover, govern, monitor, operate, and optimize backup management from one console'),
    ('How are Azure VM workloads backed up?',
     'Backup extensions (VMSnapshot for Windows, VMSnapshotLinux for Linux)'),
])

# --- AZ-104 Protect VMs with Azure Backup Assessment ---
doc.add_heading('AZ-104: Protect VMs with Azure Backup - Knowledge Check', level=2)
add_assessment_box(doc, [
    ('Which restore type replaces a disk on an existing VM?',
     'Replace existing - restores a disk and replaces it on the existing VM. VM must still exist.'),
    ('Why would replacing a disk fail?',
     'The existing VM is deleted and no longer available. Replace existing requires the VM to exist.'),
    ('What are the snapshot consistency levels?',
     'Application consistent (VSS, full memory+I/O), File system consistent (disk state, no memory), Crash consistent (disk only, VM shutdown)'),
    ('What is the difference between Azure Backup and Azure Site Recovery?',
     'Backup: Point-in-time copies for data loss/corruption/ransomware. Site Recovery: Real-time replication for failover during outages/disasters.'),
    ('What are the restore options for VMs?',
     'Create new VM, Restore disk, Replace existing, Cross Region, Cross Subscription, Cross Zonal, Selective Disk'),
    ('Can you recover individual files from a backup?',
     'Yes - Mount snapshot on target machine using iSCSI initiator to recover individual files'),
    ('What are encrypted VM restore limitations?',
     'Only standalone key encryption (not certificate keys). No file/folder restore (must restore full VM). Replace existing not available.'),
])

# --- AZ-104 Monitor Azure VMs Assessment ---
doc.add_heading('AZ-104: Monitor Azure VMs with Azure Monitor - Knowledge Checks', level=2)
add_assessment_box(doc, [
    ('What are the two main types of monitoring data?',
     'Metrics (numerical values at intervals, retained 93 days) and Logs (recorded events with timestamps, stored in Log Analytics workspace)'),
    ('What are the VM monitoring layers?',
     'VM host, Guest OS, Client workloads, and Applications'),
    ('How do you enable recommended alert rules?',
     'Select "Enable recommended alert rules" on the Monitoring tab during VM creation'),
    ('Which metric is NOT available by default on VM Monitoring tab?',
     'Guest OS Available Memory - requires VM insights and DCR setup'),
    ('How do you add another metric to Metrics Explorer graph?',
     'Select "Add metric" at upper left'),
    ('Which parameter is NOT in Metrics Explorer dropdown fields?',
     'Time range - it is set separately at upper right (not in the dropdown fields of Scope/Namespace/Metric/Aggregation)'),
    ('What does enabling VM insights provide?',
     'Prebuilt client performance workbooks, guest OS metrics collection, Azure Monitor Agent installation, preconfigured DCR'),
    ('What is the quickest way to install Azure Monitor Agent?',
     'Select Azure Monitor Agent when enabling VM insights'),
    ('How do you collect event log data from VMs?',
     'Create a Data Collection Rule (DCR)'),
    ('How do you view log data collected by a DCR?',
     'Using a KQL query in your Log Analytics workspace'),
    ('What is a DCR (Data Collection Rule)?',
     'Defines what data to collect from VMs and where to send it. Supports performance counters and event logs.'),
])

# --- AZ-104 VM Availability Assessment (9 questions) ---
doc.add_heading('AZ-104: Configure VM Availability - Assessment', level=2)
add_assessment_box(doc, [
    ('Which feature automatically adjusts VM count based on demand?',
     'Azure Virtual Machine Scale Sets with autoscale'),
    ('How to optimize cost during weekends with Scale Sets?',
     'Configure autoscale rules to decrease instances during weekends (schedule-based scaling)'),
    ('Key advantage of VM Scale Sets for demand fluctuations?',
     'Automatic adjustment of virtual machine count as demand changes'),
    ('Which scaling is best for handling increasing simultaneous user requests?',
     'Horizontal scaling (scale out) - adds more VM instances for concurrent requests'),
    ('How do update domains maintain availability during maintenance?',
     'Update domains ensure only a subset of VMs are rebooted during maintenance, minimizing downtime (one UD at a time)'),
    ('Best scaling for growing startup with low costs?',
     'Horizontal scaling with spot instances - cost-effective gradual capacity increase'),
    ('Which feature protects against datacenter failures?',
     'Availability Zones - physically separate locations within a region with independent power, cooling, networking'),
    ('How to scale database reads without impacting writes?',
     'Horizontal scaling by adding read replicas'),
    ('Which feature distributes VMs across multiple datacenters?',
     'Availability Zones - provides protection against datacenter failures'),
])

# --- AZ-104 App Service Plans Assessment ---
doc.add_heading('AZ-104: Configure App Service Plans - Assessment', level=2)
add_assessment_box(doc, [
    ('What scaling option provides more CPU, memory, disk space without adding VMs?',
     'Scale up (vertical scaling) - increase VM size/tier for more resources'),
    ('Which App Service Plan supports 10 staging slots?',
     'Premium V3 P1V3 - supports up to 20 staging slots. Standard only supports 5. Basic supports 0.'),
    ('Triggering an event at 8:00 AM on Saturday is what type of rule?',
     'A time-based rule (schedule-based autoscale rule)'),
])

# --- AZ-104 Configure App Service Assessment ---
doc.add_heading('AZ-104: Configure Azure App Service - Assessment', level=2)
add_assessment_box(doc, [
    ('When cloning a configuration from another deployment slot, which setting follows the content?',
     'Connection strings (by default swap with content unless marked as slot-specific)'),
    ('How to monitor web page usage, popularity, and user locations?',
     'Azure Application Insights - tracks page views, user locations, usage times, and performance'),
    ('Which is a valid automated deployment source?',
     'GitHub - also supports Azure DevOps, Bitbucket, Local Git, Azure Repos'),
    ('What settings are slot-specific (stay with slot)?',
     'Custom domain names, TLS/SSL settings, Scale settings, Always On, IP restrictions, CORS, Virtual network integration, Managed identities'),
    ('What settings are swapped (follow content)?',
     'Language stack/version, App settings*, Connection strings*, Public certificates, WebJobs content, Path mapping (* can be configured as slot-specific)'),
])

# --- AZ-104 Container Instances Assessment (9 questions) ---
doc.add_heading('AZ-104: Configure Container Instances - Assessment', level=2)
add_assessment_box(doc, [
    ('Key difference between ACI and VMs?',
     'Azure Container Instances do NOT require managing the underlying virtual machines'),
    ('What is a container group?',
     'A collection of containers that share a lifecycle, resources, local network, and storage volumes (similar to Kubernetes pod)'),
    ('Which feature enables internet access with a domain name for containers?',
     'Public IP connectivity and DNS names'),
    ('What could cause a multi-container group to not function?',
     'Resource allocation for the container group is insufficient (ACI adds all container resource requests together)'),
    ('Best technology for quick, isolated apps with minimal resources?',
     'Azure Container Instances - fast deployment, no VM management'),
    ('Which container advantage leads to better resource utilization?',
     'Containers use fewer system resources by running only the necessary user-mode components'),
    ('Which method deploys multiple containers?',
     'Using an orchestrator such as Azure Kubernetes Service'),
    ('Primary benefit of using ACI?',
     'Ability to run containers without managing underlying virtual machines'),
    ('Best technology for strong isolation between applications (security)?',
     'Virtual Machines - provide complete isolation with strongest security boundary'),
])

# --- AZ-104 Introduction to Azure VMs Assessment ---
doc.add_heading('AZ-104: Introduction to Azure VMs - Assessment', level=2)
add_assessment_box(doc, [
    ('Which workload option for a network appliance?',
     'Compute optimized - high CPU-to-memory ratio, suitable for network appliances, batch processes'),
    ('Are Resource Manager templates JSON files?',
     'True - ARM templates are JSON files that define resources to deploy'),
])

# --- Educator Module Assessments (for XP completion) ---
doc.add_heading('Educator Module Assessments (For XP Only)', level=2)
p = doc.add_paragraph()
run = p.add_run('Note: These educator modules are NOT part of AZ-104 or AZ-305 exams. Included only for module completion on Microsoft Learn.')
run.italic = True
run.font.size = Pt(10)

add_assessment_box(doc, [
    ('Tips & Tricks: How to prepare for complex lab environments?',
     'Begin with a demonstration of how to set up the lab environment'),
    ('Tips & Tricks: Best visual representation tool?',
     'Microsoft Whiteboard'),
    ('Tips & Tricks: Best tool for building strong foundation?',
     'Instructor-led lab experiences'),
    ('MSLE Introduction: What is MSLE?',
     'Free curriculum, training, and tools for teaching Microsoft technical skills'),
    ('MSLE Introduction: Educator resources?',
     'Access to a global community of technical Educators through the MSLE Educator Teams-based community'),
    ('MSLE Introduction: How to sign up?',
     'Sign-up at aka.ms/MSLE'),
    ('MSLE Onboard: Best place for profile, courses, vouchers, LDC?',
     'Microsoft Learn for Educators portal'),
    ('MSLE Onboard: What is in the LDC?',
     'Assessment guides (plus curriculum and teaching materials)'),
    ('MSLE Onboard: Where for blog posts and educator Q&A?',
     'Microsoft Learn for Educators Teams community'),
    ('Course Preparation: Training options for Fundamentals?',
     'Microsoft Learn and Microsoft Virtual Training Days'),
    ('Course Preparation: Where to request practice tests/vouchers?',
     'Microsoft Learn for Educators portal'),
    ('Course Preparation: How to take advanced certification exam?',
     'PearsonVUE Testing Center or Online Proctoring'),
    ('Course Planning: Way to integrate certifications?',
     'Course-level approach'),
    ('Course Planning: Tool for integrating Learn modules into curricula?',
     'LTI application'),
    ('Course Planning: Where to record course details?',
     'MSLE portal'),
    ('Best Practices: What are diverse needs addressed?',
     'Hearing, mobility, vision, and neurodiversity'),
    ('Best Practices: Goal of Microsoft ARB courses?',
     'Teach students not just the "what" but "how" to do things in Azure for certification'),
    ('Best Practices: Typical ARB exam feature?',
     'Case studies and labs'),
    ('Course Delivery: Steps to prepare for teaching?',
     'Draft syllabus, record class in MSLE portal, set up lab seats, review VEPS. Then engage community, finalize certification plan, finalize syllabus, set expectations, set up LMS.'),
    ('Course Delivery: Why are surveys important?',
     'Get insights into students experiences, feedback, and provide best tools for the course'),
])

# ============================================================
# ADDITIONAL DETAILED CONTENT - MISSING ITEMS
# ============================================================
section_break(doc)
doc.add_heading('ADDITIONAL DETAILED NOTES', level=1)
doc.add_paragraph()

# --- Azure Policy Detailed Structure ---
doc.add_heading('Azure Policy - Detailed JSON Structure', level=2)
p = doc.add_paragraph()
p.add_run('Policy Definition JSON Anatomy:').bold = True
add_bullet_list(doc, [
    'displayName: Human-readable name',
    'description: What the policy does',
    'mode: "All" (all resource types) or "Indexed" (only types supporting tags and location)',
    'metadata: Category, version, preview status',
    'parameters: Make policies reusable (e.g., allowedLocations parameter)',
    'policyRule: Contains IF (condition) and THEN (effect)',
    'Logical operators: not, allOf (AND), anyOf (OR)',
])

doc.add_heading('Policy Evaluation Triggers', level=3)
add_bullet_list(doc, [
    'Resource created, updated, or deleted',
    'Policy or initiative newly assigned to a scope',
    'Policy or initiative already assigned is updated',
    'Standard compliance evaluation cycle (every 24 hours)',
    'On-demand evaluation via REST API',
])

doc.add_heading('Compliance States', level=3)
add_table(doc, ['State', 'Meaning'], [
    ['Compliant', 'Resource meets all applicable policy conditions'],
    ['Non-compliant', 'Resource violates one or more policy conditions'],
    ['Exempt', 'Resource is exempted from evaluation (waiver or mitigated)'],
    ['Conflicting', 'Two or more policies conflict for the resource'],
    ['Not started', 'Evaluation cycle has not run yet'],
    ['Not registered', 'Resource provider not registered for policy evaluation'],
])

doc.add_heading('Safe Deployment Rings', level=3)
add_bullet_list(doc, [
    'Ring 0: DoNotEnforce mode (audit only) on small test resource group',
    'Ring 1: DoNotEnforce mode on broader scope (subscription)',
    'Ring 2: Default enforcement on test/staging subscription',
    'Ring 3: Default enforcement on production with remediation tasks',
])

# --- SSPR Additional Notes (supplements main module section) ---
doc.add_heading('Microsoft Entra SSPR - Additional Notes', level=2)
add_bullet_list(doc, [
    'SSPR Localization: Portal automatically detects browser locale and renders in appropriate language',
    'SSPR + MFA combined registration simplifies user experience (one registration for both)',
    'Security questions: Can set minimum questions to register (3-5) and minimum to reset (3-5)',
    'Security questions are stored encrypted and can only be read during SSPR flow',
    'On-premises integration: SSPR changes can trigger account unlock + password writeback simultaneously',
    'Cleanup after testing: Delete test users, delete test groups, disable SSPR, remove branding',
])

# --- Azure File Sync Detailed ---
doc.add_heading('Azure File Sync - Detailed Components', level=2)
add_table(doc, ['Component', 'Description', 'Limits'], [
    ['Storage Sync Service', 'Top-level Azure resource. Peer of storage account.', '100 sync services per subscription'],
    ['Sync Group', 'Defines sync topology for a set of files', '100 sync groups per Storage Sync Service'],
    ['Cloud Endpoint', 'Azure file share in sync group', '1 per sync group'],
    ['Server Endpoint', 'Specific path on registered Windows server', '100 per sync group'],
    ['Agent', 'Software installed on Windows Server', 'Auto-updates, supports Windows Server 2012 R2+'],
])

doc.add_heading('Cloud Tiering', level=3)
add_bullet_list(doc, [
    'Frequently accessed files cached locally on server',
    'Infrequently accessed files tiered (replaced with pointer) to Azure file share',
    'When user opens tiered file, it is seamlessly recalled from Azure',
    'Policy based: Volume free space policy (maintain X% free) and Date policy (tier files not accessed in X days)',
    'Namespace is always fully cached (all file/folder names available offline)',
])

# --- Conditional Access Detailed ---
doc.add_heading('Conditional Access - Detailed Signal Types & Decisions', level=2)

doc.add_heading('Signal Types (IF conditions)', level=3)
add_table(doc, ['Signal', 'Examples'], [
    ['User or group', 'Specific users, groups, or roles'],
    ['Cloud application', 'Target applications (Office 365, Azure portal, custom apps)'],
    ['IP location', 'Named locations - trusted IP ranges or countries'],
    ['Device platform', 'iOS, Android, Windows, macOS, Linux'],
    ['Device state', 'Hybrid joined, compliant, managed'],
    ['Client app', 'Browser, mobile apps, desktop clients, Exchange ActiveSync'],
    ['Sign-in risk', 'Real-time risk level (low/medium/high) from Identity Protection'],
    ['User risk', 'User-level risk from Identity Protection'],
])

doc.add_heading('Decisions (THEN actions)', level=3)
add_table(doc, ['Decision', 'Description'], [
    ['Block access', 'Completely prevent access'],
    ['Grant access', 'Allow with optional requirements'],
    ['Require MFA', 'Must complete multi-factor authentication'],
    ['Require compliant device', 'Device must meet Intune compliance policies'],
    ['Require Hybrid Entra Joined', 'Device must be joined to both on-premises AD and Entra ID'],
    ['Require approved client app', 'Must use approved application'],
    ['Require app protection policy', 'Intune app protection policy required'],
    ['Require password change', 'User must change password'],
    ['Session controls', 'App enforced restrictions, Conditional Access App Control, sign-in frequency, persistent browser session'],
])

# --- Storage Account Custom Domain Mapping ---
doc.add_heading('Storage Account Custom Domain Mapping', level=2)
add_table(doc, ['Method', 'Description', 'Downtime'], [
    ['Direct CNAME mapping', 'Create CNAME: blob.contoso.com > contoso.blob.core.windows.net', 'Brief downtime during DNS switch'],
    ['Intermediary mapping (asverify)', 'Create CNAME: asverify.blob.contoso.com > asverify.contoso.blob.core.windows.net, then map', 'No downtime'],
])

# --- VM Extensions Detailed ---
doc.add_heading('Azure VM Extensions - Common Types', level=2)
add_table(doc, ['Extension', 'Purpose'], [
    ['Custom Script Extension', 'Download and run scripts on VMs. Post-deployment configuration.'],
    ['DSC Extension', 'PowerShell Desired State Configuration management'],
    ['VMSnapshot / VMSnapshotLinux', 'Azure Backup snapshot extensions'],
    ['Azure Monitor Agent', 'Collect monitoring data from guest OS'],
    ['Microsoft Antimalware', 'Real-time protection and scheduled scanning'],
    ['Azure Disk Encryption', 'BitLocker (Windows) or DM-Crypt (Linux) for disk encryption'],
])

# --- Azure Automation Services ---
doc.add_heading('Azure Automation Services', level=2)
add_table(doc, ['Service', 'Description'], [
    ['Process Automation', 'Set up watcher tasks that respond to events. Automate frequent tasks with runbooks.'],
    ['Configuration Management', 'Track software updates, manage configurations. Microsoft Endpoint Configuration Manager.'],
    ['Update Management', 'Assess available updates, schedule installation, review deployment results for VMs.'],
])

doc.add_heading('Auto-shutdown', level=3)
add_bullet_list(doc, [
    'Automatically shut down VMs on a schedule to save costs',
    'Configure daily or weekly shutdown times',
    'Specify time zone for the schedule',
    'Found under VM > Operations > Auto-shutdown in Azure portal',
])

# --- Detailed VM Billing ---
doc.add_heading('VM Billing - Detailed Breakdown', level=2)
add_table(doc, ['Resource', 'Cost'], [
    ['Virtual Network', 'See Virtual Network pricing'],
    ['NIC (Network Interface Card)', 'No separate cost (limit based on VM size)'],
    ['Private IP / Public IP', 'See IP Addresses pricing'],
    ['NSG', 'No additional charges'],
    ['OS Disk + Data Disks', 'Charged at Managed Disks rates'],
    ['Local/Temporary Disk', 'Free (no charge for local disk storage)'],
    ['OS License (Windows)', 'Included in hourly rate (reduce with Azure Hybrid Benefit)'],
    ['OS License (Linux)', 'No OS license charge (cheaper than Windows)'],
])

add_key_point(doc, 'Compute: Billed per-minute. Stop+deallocate = no compute charge. Stop only = still charged.')
add_key_point(doc, 'Storage: Charged even when VM is deallocated. Must delete disks to stop storage charges.')
add_key_point(doc, 'Reserved Instances: Up to 72% savings with 1 or 3 year commitment')

# --- Azure Container Apps Detailed ---
doc.add_heading('Azure Container Apps - Detailed Features', level=2)
add_bullet_list(doc, [
    'Serverless platform for containerized applications',
    'Powered by Kubernetes, Dapr, KEDA, and envoy',
    'Scale based on: HTTP traffic, event-driven processing, CPU/memory load, any KEDA-supported scaler',
    'Supports scale to zero (no charges when idle)',
    'Service discovery and traffic splitting for microservices',
    'No direct access to underlying Kubernetes APIs (use AKS if needed)',
    'Common uses: API endpoints, background processing jobs, event-driven processing, microservices',
])

doc.add_heading('ACA vs AKS Comparison', level=3)
add_table(doc, ['Feature', 'Container Apps (ACA)', 'Kubernetes Service (AKS)'], [
    ['Management', 'Simplified PaaS, built on AKS', 'Full Kubernetes control'],
    ['Deployment', 'Quick, minimal configuration', 'More setup, more customization'],
    ['Scaling', 'HTTP-based + event-driven autoscaling', 'Horizontal pod + cluster autoscaling'],
    ['Expertise needed', 'Low (no K8s knowledge required)', 'High (Kubernetes expertise needed)'],
    ['Best for', 'Microservices, serverless workloads', 'Complex, long-running enterprise apps'],
])

# --- Azure Monitor Detailed ---
doc.add_heading('Azure Monitor - Additional Details', level=2)

doc.add_heading('Recommended Alert Rules (VM Host)', level=3)
add_table(doc, ['Alert', 'Description'], [
    ['VM Availability', 'Alerts when VM stops running'],
    ['CPU Percentage', 'High CPU utilization threshold'],
    ['Available Memory', 'Low available memory threshold'],
    ['OS Disk IOPS Consumed', 'Disk performance threshold'],
    ['Network In/Out Total', 'Network traffic thresholds'],
    ['Data Disk Latency', 'Disk latency threshold'],
])

doc.add_heading('Boot Diagnostics', level=3)
add_bullet_list(doc, [
    'Host logs for troubleshooting boot issues',
    'Enable during VM creation or afterwards for existing VMs',
    'View screenshots from VM hypervisor (Windows and Linux)',
    'View serial console log output of boot sequence (Linux)',
    'Data stored in managed storage account',
])

doc.add_heading('Data Collection Endpoints', level=3)
add_bullet_list(doc, [
    'Required for sending log data to Azure Monitor',
    'Created in Azure Monitor > Settings > Data Collection Endpoints',
    'Must be in same region as VM',
    'Associated with DCRs for log collection',
])

# ============================================================
# EXAM TIPS & KEY FACTS TO REMEMBER
# ============================================================
section_break(doc)
doc.add_heading('EXAM TIPS: Key Facts to Remember', level=1)
doc.add_paragraph()

doc.add_heading('Numbers to Remember', level=2)
add_table(doc, ['Item', 'Number'], [
    ['Management group depth', '6 levels (excl. root and subscription)'],
    ['Reserved IPs per subnet', '5'],
    ['NSG rule priority range', '100-4096'],
    ['Default update domains', '5 (max 20)'],
    ['Default fault domains', '2-3'],
    ['Availability zones per region', '3'],
    ['Max tags per resource', '50'],
    ['Tag name max length', '512 characters'],
    ['Tag value max length', '256 characters'],
    ['ARM template max parameters', '256'],
    ['VM name max (Windows)', '15 characters'],
    ['VM name max (Linux)', '64 characters'],
    ['File share snapshots max', '200 per share'],
    ['Soft delete retention for backup', '14 days (free)'],
    ['Soft delete for files retention', '1-365 days'],
    ['Azure Monitor metrics retention', '93 days'],
    ['Cloud Shell idle timeout', '20 minutes'],
    ['Policy evaluation cycle', 'Every 24 hours'],
    ['Sync groups per Storage Sync Service', '100'],
    ['Server endpoints per sync group', '100'],
    ['Blob max size (Block)', '190.7 TB'],
    ['Blob max size (Page)', '8 TB'],
    ['Cool tier min retention', '30 days'],
    ['Cold tier min retention', '90 days'],
    ['Archive tier min retention', '180 days'],
    ['Scale Sets max instances', '1,000 (600 custom image)'],
    ['App Service Free staging slots', '0'],
    ['App Service Standard staging slots', '5'],
    ['App Service Premium staging slots', '20'],
    ['App Service Isolated max instances', '200'],
])

doc.add_heading('Key "CANNOT" Facts', level=2)
add_bullet_list(doc, [
    'Resource groups CANNOT be nested',
    'Resource groups CANNOT be renamed',
    'Tags are NOT inherited from resource group to resources',
    'VNets CANNOT be shared across subscriptions',
    'VNet address spaces CANNOT overlap when peering',
    'VNet peering is NOT transitive',
    'CNAME records CANNOT be used at zone apex (use Alias records)',
    'NSG default rules CANNOT be deleted (only overridden)',
    'Security questions CANNOT be used by admins for SSPR',
    'Ultra Disk and Premium SSD v2 CANNOT be used as OS disks',
    'Archive tier is OFFLINE - must rehydrate before access',
    'Blob snapshots NOT supported for object replication',
    'Port mapping NOT supported in container groups',
    'Replace existing restore NOT available for encrypted VMs or deleted VMs',
    'Network Watcher is for IaaS ONLY - not PaaS',
    'Azure DNS does NOT support domain name purchasing',
])

doc.add_heading('Key "ALWAYS/DEFAULT" Facts', level=2)
add_bullet_list(doc, [
    'Standard SKU public IP: ALWAYS static allocation',
    'Storage encryption (SSE): ALWAYS enabled, cannot be disabled',
    'HTTPS: ALWAYS enforced by default for storage',
    'Admin accounts: ALWAYS require 2 authentication methods for SSPR',
    'Default update domains: 5 (immutable after creation)',
    'Default NSG inbound: DenyAllInBound',
    'Default NSG outbound: AllowInternetOutBound',
    'Inbound NSG order: Subnet NSG FIRST, then NIC NSG',
    'Outbound NSG order: NIC NSG FIRST, then Subnet NSG',
    'Load Balancer default: 5-tuple hash (no session persistence)',
    'Blob container default access: Private (no anonymous access)',
    'Network Watcher: Automatically available when VNet is created',
    'Soft delete: Enabled by default for Key Vault',
    'Azure Backup soft delete: 14 days free retention',
])

# ============================================================
# SAVE DOCUMENT
# ============================================================
print("Saving document...")
doc.save('E:/Azure/Azure_Certification_Study_Guide_AZ104_AZ305.docx')
print("Document saved successfully!")
print(f"Location: E:/Azure/Azure_Certification_Study_Guide_AZ104_AZ305.docx")
